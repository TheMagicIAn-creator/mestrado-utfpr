"""
leitor_anexos.py — Al IAdo PV

Leitura ad-hoc de arquivos anexados na conversa — o equivalente, para o Al IAdo
PV, ao "anexar um arquivo" no ChatGPT/Claude. NAO indexa nada permanentemente
(isso e papel do watcher + processador_pdf); apenas extrai o conteudo do anexo
para que o LLM possa ler, interpretar e consultar na resposta ATUAL.

Tipos suportados:
  - PDF                         → texto (pypdf), capado
  - CSV / TSV                   → preview tabular (pandas)
  - XLSX / XLS                  → preview por aba (pandas + openpyxl)
  - DOCX                        → paragrafos + tabelas (python-docx)
  - TXT/MD/codigo/JSON/YAML/... → decode de texto, capado
  - PNG/JPG/JPEG/GIF/WEBP/BMP   → base64 (para Gemini multimodal)
  - desconhecido                → tenta texto; se binario, nota descritiva

A funcao publica `ler_anexos([(nome, bytes), ...])` devolve uma lista de dicts
normalizados; `montar_bloco_texto_anexos(...)` consolida o texto para o prompt.

Autor: Rodolfo Torres (UTFPR)
"""

from __future__ import annotations

import base64
import io
from pathlib import Path


# ============================================================
# LIMITES (protegem o orcamento do prompt e os limites do LLM)
# ============================================================

MAX_CHARS_POR_ANEXO = 12_000   # texto extraido de UM anexo
MAX_CHARS_TOTAL     = 30_000   # soma de texto de TODOS os anexos
MAX_IMAGENS         = 4        # imagens por mensagem
MAX_DIM_IMAGEM      = 1_024    # px no maior lado (downscale)
MAX_LINHAS_PREVIEW  = 40       # linhas de preview de CSV/Excel


EXTS_IMAGEM = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}

EXTS_TEXTO = {
    ".txt", ".md", ".markdown", ".rst", ".text",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".c", ".cpp", ".cc",
    ".h", ".hpp", ".cs", ".go", ".rs", ".rb", ".php", ".swift", ".kt",
    ".scala", ".r", ".m", ".jl", ".lua", ".pl",
    ".sql", ".sh", ".bash", ".bat", ".ps1",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".env",
    ".html", ".htm", ".xml", ".css", ".scss", ".tex", ".log", ".csv2",
}


# ============================================================
# AUXILIARES
# ============================================================

def _cap(texto: str | None, limite: int = MAX_CHARS_POR_ANEXO) -> str:
    """Trunca texto preservando o inicio e sinalizando o corte."""
    if not texto:
        return ""
    if len(texto) <= limite:
        return texto
    return texto[:limite].rstrip() + f"\n\n[...conteudo truncado em {limite} caracteres...]"


def _decode(dados: bytes) -> str:
    """Decodifica bytes tentando varias codificacoes; nunca explode."""
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return dados.decode(enc)
        except Exception:
            continue
    return dados.decode("utf-8", errors="replace")


def _parece_binario(dados: bytes) -> bool:
    """Heuristica simples: bytes nulos nos primeiros KB indicam binario."""
    return b"\x00" in dados[:4096]


def _df_para_texto(df, nome_aba: str | None = None) -> str:
    """Representacao textual compacta de um DataFrame para o LLM."""
    cab = f"### Aba: {nome_aba}\n" if nome_aba else ""
    info = f"{df.shape[0]} linhas x {df.shape[1]} colunas\n"
    colunas = "Colunas: " + ", ".join(str(c) for c in df.columns) + "\n\n"
    amostra = df.head(MAX_LINHAS_PREVIEW)
    try:
        preview = amostra.to_markdown(index=False)   # requer tabulate
    except Exception:
        preview = amostra.to_string(index=False)
    return cab + info + colunas + preview


# ============================================================
# EXTRATORES POR TIPO
# ============================================================

def _ler_pdf(dados: bytes) -> tuple[str, str]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(dados))
    partes: list[str] = []
    total = 0
    for pagina in reader.pages:
        try:
            t = pagina.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            partes.append(t)
            total += len(t)
        if total > MAX_CHARS_POR_ANEXO:
            break
    texto = "\n".join(partes).strip()
    return texto, f"PDF, {len(reader.pages)} pagina(s)"


def _ler_csv(dados: bytes, sep: str | None = None) -> tuple[str, str]:
    import pandas as pd

    try:
        if sep is None:
            df = pd.read_csv(io.BytesIO(dados), sep=None, engine="python")
        else:
            df = pd.read_csv(io.BytesIO(dados), sep=sep)
    except Exception:
        return _decode(dados), "tabela (lida como texto bruto)"
    return _df_para_texto(df), f"tabela {df.shape[0]} linhas x {df.shape[1]} colunas"


def _ler_excel(dados: bytes) -> tuple[str, str]:
    import pandas as pd

    try:
        abas = pd.read_excel(io.BytesIO(dados), sheet_name=None)  # requer openpyxl
    except Exception as e:
        return "", f"Excel ilegivel ({type(e).__name__})"
    partes: list[str] = []
    total = 0
    for nome_aba, df in abas.items():
        bloco = _df_para_texto(df, nome_aba=str(nome_aba))
        partes.append(bloco)
        total += len(bloco)
        if total > MAX_CHARS_POR_ANEXO:
            break
    return "\n\n".join(partes).strip(), f"planilha Excel, {len(abas)} aba(s)"


def _ler_docx(dados: bytes) -> tuple[str, str]:
    import docx  # python-docx

    documento = docx.Document(io.BytesIO(dados))
    linhas = [p.text for p in documento.paragraphs if p.text.strip()]
    for tabela in documento.tables:
        for row in tabela.rows:
            celulas = [c.text.strip() for c in row.cells]
            if any(celulas):
                linhas.append(" | ".join(celulas))
    return "\n".join(linhas).strip(), f"documento Word, {len(documento.paragraphs)} paragrafo(s)"


def _ler_imagem(dados: bytes, ext: str) -> tuple[str, str, str]:
    from PIL import Image

    img = Image.open(io.BytesIO(dados))
    img.load()
    formato = (img.format or ext.lstrip(".")).upper()
    largura, altura = img.size

    if max(img.size) > MAX_DIM_IMAGEM:
        ratio = MAX_DIM_IMAGEM / float(max(img.size))
        img = img.resize((max(1, int(img.width * ratio)), max(1, int(img.height * ratio))))

    out = io.BytesIO()
    if img.mode in ("RGBA", "P", "LA"):
        img.convert("RGBA").save(out, format="PNG")
        mime = "image/png"
    else:
        img.convert("RGB").save(out, format="JPEG", quality=85)
        mime = "image/jpeg"

    b64 = base64.b64encode(out.getvalue()).decode("ascii")
    return b64, mime, f"imagem {formato} {largura}x{altura}px"


# ============================================================
# DISPATCHER
# ============================================================

def ler_anexo(nome: str, dados: bytes) -> dict:
    """
    Le UM anexo (nome + bytes) e devolve um dict normalizado:
      {nome, tipo, texto, imagem_b64, mime, resumo, erro}
    tipo ∈ {"texto", "imagem", "erro"}.
    """
    resultado = {
        "nome": nome,
        "tipo": "erro",
        "texto": "",
        "imagem_b64": "",
        "mime": "",
        "resumo": "",
        "erro": None,
    }

    try:
        ext = Path(nome).suffix.lower()
        tamanho = len(dados or b"")

        if not dados:
            resultado["erro"] = "arquivo vazio"
            resultado["resumo"] = f"{ext or '(sem extensao)'} vazio"
            return resultado

        # ── Imagem ───────────────────────────────────────────
        if ext in EXTS_IMAGEM:
            b64, mime, resumo = _ler_imagem(dados, ext)
            resultado.update(tipo="imagem", imagem_b64=b64, mime=mime, resumo=resumo)
            return resultado

        # ── Documentos com extrator dedicado ─────────────────
        if ext == ".pdf":
            texto, resumo = _ler_pdf(dados)
        elif ext in (".csv",):
            texto, resumo = _ler_csv(dados, sep=None)
        elif ext in (".tsv",):
            texto, resumo = _ler_csv(dados, sep="\t")
        elif ext in (".xlsx", ".xls", ".xlsm"):
            texto, resumo = _ler_excel(dados)
        elif ext == ".docx":
            texto, resumo = _ler_docx(dados)
        elif ext in EXTS_TEXTO:
            texto, resumo = _decode(dados), f"texto {ext} ({tamanho} bytes)"
        else:
            # Desconhecido: se binario, nota descritiva; senao trata como texto.
            if _parece_binario(dados):
                resultado.update(
                    tipo="erro",
                    erro=f"tipo nao suportado ({ext or 'sem extensao'}); binario de {tamanho} bytes",
                    resumo=f"arquivo binario {ext or '(sem extensao)'} ({tamanho} bytes)",
                )
                return resultado
            texto, resumo = _decode(dados), f"texto {ext or '(sem extensao)'} ({tamanho} bytes)"

        texto = _cap(texto)
        if not texto.strip():
            resultado.update(tipo="erro", erro="sem texto extraivel", resumo=resumo)
            return resultado

        resultado.update(tipo="texto", texto=texto, resumo=resumo)
        return resultado

    except Exception as e:  # noqa: BLE001 — qualquer falha vira "erro" amigavel
        resultado["erro"] = f"{type(e).__name__}: {e}"
        return resultado


def ler_anexos(arquivos: list[tuple[str, bytes]]) -> list[dict]:
    """
    Le uma lista de (nome, bytes), aplicando os limites globais de texto e de
    numero de imagens. Retorna a lista de dicts normalizados.
    """
    saida: list[dict] = []
    total_chars = 0
    n_imagens = 0

    for nome, dados in (arquivos or []):
        r = ler_anexo(nome, dados)

        if r["tipo"] == "texto":
            if total_chars >= MAX_CHARS_TOTAL:
                r["texto"] = "[anexo omitido: limite total de texto dos anexos atingido]"
            elif total_chars + len(r["texto"]) > MAX_CHARS_TOTAL:
                restante = MAX_CHARS_TOTAL - total_chars
                r["texto"] = _cap(r["texto"], restante)
            total_chars += len(r["texto"])

        elif r["tipo"] == "imagem":
            n_imagens += 1
            if n_imagens > MAX_IMAGENS:
                r.update(
                    tipo="erro",
                    imagem_b64="",
                    erro=f"limite de {MAX_IMAGENS} imagens por mensagem atingido",
                )

        saida.append(r)

    return saida


# ============================================================
# MONTAGEM PARA O PROMPT
# ============================================================

def montar_bloco_texto_anexos(anexos: list[dict], suporta_imagem: bool = True) -> str:
    """
    Consolida os anexos em texto para o prompt do LLM. Anexos de texto entram
    com seu conteudo; imagens entram como nota (o pixel vai pela via multimodal
    quando suporta_imagem=True) e erros como aviso curto.
    """
    if not anexos:
        return ""

    partes: list[str] = []
    for a in anexos:
        nome = a.get("nome", "anexo")
        tipo = a.get("tipo")
        resumo = a.get("resumo", "")

        if tipo == "texto":
            partes.append(f"[Anexo: {nome} — {resumo}]\n{a.get('texto', '')}")
        elif tipo == "imagem":
            if suporta_imagem:
                partes.append(
                    f"[Anexo (imagem): {nome} — {resumo}. "
                    "A imagem segue anexada para analise visual; descreva e use o que for pertinente.]"
                )
            else:
                partes.append(
                    f"[Anexo (imagem): {nome} — {resumo}. "
                    "O provedor de LLM atual nao le imagens (texto puro). "
                    "Avise o Rodolfo que, para analise visual, ele deve conectar o Google Gemini.]"
                )
        else:
            partes.append(f"[Anexo nao lido: {nome} — {a.get('erro') or 'erro desconhecido'}]")

    return "\n\n".join(partes)


def tem_imagem(anexos: list[dict]) -> bool:
    """True se algum anexo e imagem com bytes utilizaveis."""
    return any(a.get("tipo") == "imagem" and a.get("imagem_b64") for a in (anexos or []))
