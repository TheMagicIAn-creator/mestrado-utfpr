from __future__ import annotations

import io
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable


RAIZ = Path(__file__).resolve().parents[1]
if str(RAIZ) not in sys.path:
    sys.path.insert(0, str(RAIZ))

import chromadb
from sentence_transformers import SentenceTransformer

from src.conhecimento.agente import (
    _expandir_query,
    _montar_prompt,
    buscar_contexto,
    catalogo_literatura,
    deve_consultar_literatura,
    eh_query_de_revisao,
    formatar_referencias_markdown,
    remover_bloco_fontes_llm,
    resposta_interacao_simples,
)
from src.conhecimento.ferramentas import decidir_acao, executar_ferramenta
from src.conhecimento.leitor_anexos import (
    ler_anexo,
    ler_anexos,
    montar_bloco_texto_anexos,
)
from src.conhecimento.indexador import upsert_em_lotes
from src.core.config import (
    MODELO_EMBEDDINGS,
    NOME_COLECAO,
    NOME_COLECAO_AVALIACOES,
    NOME_COLECAO_SESSOES,
    PASTA_CHROMADB,
    PASTA_NOTAS,
)

# Modelo de embeddings e colecao da literatura — carregados sob demanda
# pelos casos que checam recuperacao real.
_RAG_CACHE: dict = {}


def _rag_cache() -> tuple:
    if "modelo" not in _RAG_CACHE:
        _RAG_CACHE["modelo"] = SentenceTransformer(MODELO_EMBEDDINGS)
        client = chromadb.PersistentClient(path=str(PASTA_CHROMADB))
        _RAG_CACHE["colecao"] = client.get_or_create_collection(name=NOME_COLECAO)
    return _RAG_CACHE["modelo"], _RAG_CACHE["colecao"]


@dataclass
class CasoTeste:
    nome: str
    categoria: str
    pergunta: str
    executar: Callable[[], tuple[bool, str]]


ORCAMENTO_TESTE = {
    "contexto_chars": 2_500,
    "historico_turnos": 2,
    "historico_chars": 600,
    "max_prompt_chars": 20_000,
    "anexos_chars": 6_000,
}


def caso_literatura(nome: str, pergunta: str, esperado: bool) -> CasoTeste:
    def executar() -> tuple[bool, str]:
        obtido = deve_consultar_literatura(pergunta)
        ok = obtido is esperado
        return ok, f"esperado={esperado}; obtido={obtido}"

    categoria = "literatura_explicita" if esperado else "sem_literatura"
    return CasoTeste(nome, categoria, pergunta, executar)


def caso_ferramenta(nome: str,
                    pergunta: str,
                    ferramenta_esperada: str | None,
                    usar_esperado: bool = True) -> CasoTeste:
    def executar() -> tuple[bool, str]:
        decisao = decidir_acao(pergunta, None)
        ferramenta = decisao.get("ferramenta")
        usar = bool(decisao.get("usar_ferramenta"))
        ok = usar is usar_esperado and ferramenta == ferramenta_esperada
        return ok, f"esperado=({usar_esperado}, {ferramenta_esperada}); obtido=({usar}, {ferramenta})"

    return CasoTeste(nome, "roteamento_ferramentas", pergunta, executar)


def caso_interacao(nome: str, pergunta: str) -> CasoTeste:
    termos_proibidos = ("literatura", "referencia", "referências", "fontes", "artigos")

    def executar() -> tuple[bool, str]:
        resposta = resposta_interacao_simples(pergunta)
        texto = (resposta or "").lower()
        proibidos = [t for t in termos_proibidos if t in texto]
        ok = bool(resposta) and not proibidos
        return ok, f"resposta_local={bool(resposta)}; termos_proibidos={proibidos or 'nenhum'}"

    return CasoTeste(nome, "interacao_humana", pergunta, executar)


def caso_referencias(nome: str, entrada, esperado: str) -> CasoTeste:
    def executar() -> tuple[bool, str]:
        obtido = formatar_referencias_markdown(entrada)
        ok = obtido == esperado
        return ok, f"esperado={esperado!r}; obtido={obtido!r}"

    return CasoTeste(nome, "referencias_markdown", str(entrada), executar)


def caso_diversidade_literatura(nome: str,
                                pergunta: str,
                                min_fontes: int,
                                pastas_proibidas: tuple[str, ...] = (),
                                arquivos_proibidos: tuple[str, ...] = ()) -> CasoTeste:
    """
    Exige que a recuperacao real (busca_hibrida + rerank) traga ao menos
    `min_fontes` documentos distintos e que nenhuma fonte caia em pastas
    ou arquivos proibidos (ex.: textbooks generalistas).
    """
    def executar() -> tuple[bool, str]:
        modelo, colecao = _rag_cache()
        _ctx, citacoes = buscar_contexto(
            pergunta,
            modelo,
            colecao,
            n_pool=120,
            n_resultados=16,
            n_resultados_revisao=28,
            max_chunks_por_fonte=2,
            contexto_chars=14_000,
            sessao_chars=1_500,
            consultar_literatura=True,
        )
        fontes = list(citacoes.keys())
        n = len(fontes)
        proibidos_arq = [a for a in arquivos_proibidos if a in fontes]
        proibidos_pasta = []
        if pastas_proibidas:
            try:
                metas = colecao.get(
                    where={"arquivo": {"$in": fontes}} if fontes else None,
                    include=["metadatas"],
                ) if fontes else {"metadatas": []}
                for m in metas.get("metadatas", []):
                    pasta = m.get("pasta", "")
                    if pasta in pastas_proibidas and m.get("arquivo") not in proibidos_pasta:
                        proibidos_pasta.append(m.get("arquivo"))
            except Exception:
                pass

        ok = n >= min_fontes and not proibidos_arq and not proibidos_pasta
        detalhe = (
            f"n_fontes={n} (min={min_fontes}); "
            f"arq_proibidos={proibidos_arq or 'nenhum'}; "
            f"pasta_proibidos={proibidos_pasta or 'nenhum'}"
        )
        return ok, detalhe

    return CasoTeste(nome, "diversidade_literatura", pergunta, executar)


def caso_strip_fontes(nome: str,
                      entrada: str,
                      deve_remover: tuple[str, ...],
                      deve_preservar: tuple[str, ...]) -> CasoTeste:
    """
    Testa o helper remover_bloco_fontes_llm: cabecalhos de Referencias/
    Bibliografia/📚 Fontes devem ser cortados, mas conteudo legitimo
    (texto da resposta, citacoes inline) deve ser preservado.
    """
    def executar() -> tuple[bool, str]:
        saida = remover_bloco_fontes_llm(entrada)
        nao_removidos = [t for t in deve_remover if t in saida]
        nao_preservados = [t for t in deve_preservar if t not in saida]
        ok = not nao_removidos and not nao_preservados
        detalhe = (
            f"nao_removidos={nao_removidos or 'nenhum'}; "
            f"nao_preservados={nao_preservados or 'nenhum'}"
        )
        return ok, detalhe

    return CasoTeste(nome, "strip_fontes_llm", entrada[:60], executar)


def caso_expansao_revisao(nome: str,
                          pergunta: str,
                          eh_revisao_esperado: bool,
                          topicos_esperados: tuple[str, ...] = ()) -> CasoTeste:
    """
    Verifica o detector eh_query_de_revisao e, quando True, a expansao
    de query injeta os topicos da dissertacao.
    """
    def executar() -> tuple[bool, str]:
        eh_rev = eh_query_de_revisao(pergunta)
        expansao = _expandir_query(pergunta)
        variacoes = " || ".join(expansao.get("variacoes", []))
        if eh_rev != eh_revisao_esperado:
            return False, f"esperado_revisao={eh_revisao_esperado}; obtido={eh_rev}"
        if eh_revisao_esperado:
            faltando = [t for t in topicos_esperados if t.lower() not in variacoes.lower()]
            if faltando:
                return False, f"topicos_faltando={faltando}"
        return True, f"revisao={eh_rev}; n_variacoes={len(expansao.get('variacoes', []))}"

    return CasoTeste(nome, "expansao_revisao", pergunta, executar)


def caso_proveniencia(nome: str,
                      pergunta: str,
                      arquivos_esperados: tuple[str, ...],
                      min_match: int) -> CasoTeste:
    """
    Para uma pergunta topica especifica, exige que pelo menos `min_match`
    dos arquivos esperados (papers do dominio correto) apareca no top-K
    da recuperacao. Detecta regressoes onde o RAG perde papers core.
    """
    def executar() -> tuple[bool, str]:
        modelo, colecao = _rag_cache()
        _ctx, citacoes = buscar_contexto(
            pergunta,
            modelo,
            colecao,
            n_pool=120,
            n_resultados=16,
            n_resultados_revisao=28,
            max_chunks_por_fonte=2,
            contexto_chars=14_000,
            sessao_chars=1_500,
            consultar_literatura=True,
        )
        # chaves podem ser 'arquivo' (legado) ou 'arquivo|pag|pag|hash'
        # (citacao por pagina) — comparamos pelo arquivo-base.
        fontes = {str(k).split('|')[0] for k in citacoes}
        encontrados = [a for a in arquivos_esperados if a in fontes]
        ok = len(encontrados) >= min_match
        return ok, (
            f"encontrados={len(encontrados)}/{len(arquivos_esperados)} "
            f"(min={min_match}); papers={encontrados or 'nenhum'}"
        )

    return CasoTeste(nome, "proveniencia_topica", pergunta, executar)


def caso_contexto_diverso(nome: str,
                          pergunta: str,
                          min_fontes_no_contexto: int) -> CasoTeste:
    """
    Verifica que o CONTEXTO efetivamente montado para o LLM contem
    pelo menos N marcadores `[Fonte: ...]` distintos — i.e., o LLM
    recebe diversidade real, nao so a lista de citacoes.
    """
    import re

    def executar() -> tuple[bool, str]:
        modelo, colecao = _rag_cache()
        ctx, _ = buscar_contexto(
            pergunta,
            modelo,
            colecao,
            n_pool=120,
            n_resultados=16,
            n_resultados_revisao=28,
            max_chunks_por_fonte=2,
            contexto_chars=14_000,
            sessao_chars=1_500,
            consultar_literatura=True,
        )
        fontes_no_ctx = set(re.findall(r"\[Fonte:\s*([^\]]+)\]", ctx))
        ok = len(fontes_no_ctx) >= min_fontes_no_contexto
        return ok, (
            f"fontes_no_contexto={len(fontes_no_ctx)} "
            f"(min={min_fontes_no_contexto})"
        )

    return CasoTeste(nome, "contexto_diverso", pergunta, executar)


def caso_prompt(nome: str,
                pergunta: str,
                consultar_literatura: bool,
                deve_conter: tuple[str, ...],
                nao_deve_conter: tuple[str, ...] = ()) -> CasoTeste:
    def executar() -> tuple[bool, str]:
        prompt = _montar_prompt(
            pergunta=pergunta,
            contexto="Memoria do projeto: resultados AUC=0,935 e F1 alto.",
            historico_formatado="",
            orcamento=ORCAMENTO_TESTE,
            consultar_literatura=consultar_literatura,
        )
        faltando = [trecho for trecho in deve_conter if trecho not in prompt]
        indevidos = [trecho for trecho in nao_deve_conter if trecho in prompt]
        ok = not faltando and not indevidos
        return ok, f"faltando={faltando or 'nenhum'}; indevidos={indevidos or 'nenhum'}"

    categoria = "prompt_com_literatura" if consultar_literatura else "prompt_sem_literatura"
    return CasoTeste(nome, categoria, pergunta, executar)


# ============================================================
# GERACAO EM ESCALA — sem/com literatura por template x topico
# ============================================================
# Topicos "seguros": nenhum contem gatilho de literatura/autor, entao a
# classificacao depende exclusivamente do template (provado pelo probe).
_TOPICOS_TESTE = [
    "FMEA", "FMECA", "RCM", "Weibull", "RUL", "autoencoder",
    "isolation forest", "random forest", "XGBoost", "gradient boosting",
    "MTTF", "B10", "filtro LCL", "IGBT", "THD", "FFT", "RMS",
    "deteccao de anomalia", "manutencao preditiva", "monitoramento de condicao",
    "prognostico", "erro de reconstrucao", "limiar p99", "baseline saudavel",
    "feature engineering", "severidade minima detectavel", "curva ROC",
    "matriz de confusao", "modo de falha", "criticidade", "NPR",
    "desbalanceamento", "sensor CA", "duty cycle", "validacao cruzada",
    "overfitting", "normalizacao de dados",
]

# Templates SEM gatilho de literatura → deve_consultar_literatura == False
_SEM_TPL = [
    "Explique {T}.",
    "Como funciona {T}?",
    "Qual a intuicao de {T}?",
    "Me ajude a entender {T}.",
    "Interprete {T} no nosso contexto.",
    "Resuma {T} para a defesa.",
    "Como aplico {T} no projeto?",
    "De um exemplo pratico de {T}.",
    "Quais os passos para usar {T}?",
    "Por que {T} importa na metodologia?",
]

# Templates COM gatilho de literatura → deve_consultar_literatura == True
_COM_TPL = [
    "Cite artigos sobre {T}.",
    "Segundo a literatura, explique {T}.",
    "Com base na literatura, descreva {T}.",
    "O que a bibliografia diz sobre {T}?",
    "Liste referencias sobre {T}.",
    "Quais papers tratam de {T}?",
    "Cite autores que discutem {T}.",
    "Ha fontes sobre {T} para eu ler?",
    "Faca uma revisao bibliografica de {T}.",
    "Levante o estado da arte de {T}.",
]


def _gerar(templates: list[str], n_por_template: int) -> list[str]:
    """Gera perguntas template x topico, ciclando os topicos."""
    out: list[str] = []
    ti = 0
    for tpl in templates:
        for _ in range(n_por_template):
            out.append(tpl.format(T=_TOPICOS_TESTE[ti % len(_TOPICOS_TESTE)]))
            ti += 1
    return out


# ============================================================
# LISTA AUTORITATIVA DOS 39 DOCUMENTOS INDEXADOS
# ============================================================
# Blindagem direta da preocupacao "alem do NASA, os outros 33 aparecem?".
# Para CADA arquivo: dispara RAG pelo sobrenome (1o token do nome) e exige
# que o proprio arquivo volte nas citacoes. Mesmo formato do verificar_autores.
ARQUIVOS_INDEXADOS = [
    "administration_nasa-reliability-centered-maintenance-guide-for-facilities-a_2008.pdf",
    "ahirwar_enhanced-anomaly-detection-in-solar-power-plants-using-hybri_2025.pdf",
    "carpinetti_gestao-da-qualidade-cap-6_2016.pdf",
    "cristaldi_a-root-cause-analysis-and-a-risk-evaluation-of-pv-balance-of_2017.pdf",
    "dhople_estimation-of-photovoltaic-system-reliability-and-performanc_2012.pdf",
    "diniz_digital-signal-processing-system-analysis-and-design_2021.pdf",
    "eletrica_subestacoes-de-energia-definicoes-conceitos-e-aplicacoes_0000.pdf",
    "francisti_predictive-modeling-and-anomaly-detection-in-solar-pv-invert_2025.pdf",
    "frontin_equipamentos-de-alta-tensao-prospeccao-e-hierarquizacao_2013.pdf",
    "ghoneim_fault-detection-algorithms-for-achieving-service-continuity-_2021.pdf",
    "gonzalez_digital-image-processing_2008.pdf",
    "grewal_kalman-filtering-theory-and-practice-using-matlab_2001.pdf",
    "grewal_power-electronics-chapter-8_2002.pdf",
    "ibrahim_machine-learning-schemes-for-anomaly-detection-in-solar-powe_2022.pdf",
    "joshi_reliability-estimation-for-components-of-photovoltaic-system_1996.pdf",
    "karim_a-review-on-risk-and-reliability-analysis-in-photovoltaic-po_2025.pdf",
    "lafraia_manual-de-confiabilidade-mantenabilidade-e-disponibilidade-cap4_0000.pdf",
    "lafraia_manual-de-confiabilidade-mantenabilidade-e-disponibilidade_0000.pdf",
    "marangis_intelligent-maintenance-approaches-for-improving-photovoltai_2025.pdf",
    "monteiro_identifying-critical-failures-in-pv-systems-based-on-pv-inve_2024.pdf",
    "moura_engenharia-de-sistemas-de-potencia-transmissao-de-energia-el_2019.pdf",
    "muqauwim_analysis-of-optimal-maintenance-interval-on-id-fan-using-rel_2020.pdf",
    "narayanan_machine-learning-based-explainable-fault-detection-of-vacuum_2023.pdf",
    "oppenheim_discrete-time-signal-processing_2014.pdf",
    "pahwa_design-and-estimation-of-reliability-of-an-off-grid-solar-ph_2017.pdf",
    "patil_a-reliability-and-risk-assessment-of-solar-photovoltaic-pane_2024.pdf",
    "puc-rio_analise-da-confiabilidade-em-sistemas-de-potencia_2003.pdf",
    "risi_advancing-solar-pv-component-inspection-early-defect-detecti_2023.pdf",
    "sakurada_as-tecnicas-de-analise-do-modos-de-falhas-e-seus-efeitos-e-a_1998.pdf",
    "sharma_a-self-tuning-reinforcement-learning-driven-isolation-forest_2026.pdf",
    "shuttleworth_reliability-prediction-of-pv-inverters-based-on-mil-hdbk-217_2015.pdf",
    "silva_avaliacao-da-confiabilidade-em-sistemas-eletricos-com-base-n_2008.pdf",
    "smith_the-scientist-and-engineer-s-guide-to-digital-signal-process_1999.pdf",
    "stender_data-set-description-three-phase-igbt-two-level-inverter-for_2020.pdf",
    "stewart_calculo-volume-i_2013.pdf",
    "tekalp_digital-video-processing_2015.pdf",
    "torres_aplicacao-da-metodologia-reliability-centred-maintenance-a-s_2024.pdf",
    "voss_service-service-architecture-yield-monitoring-optimization-a_2009.pdf",
    "xavier_analise-de-confiabilidade-em-sistemas-de-potencia_2005.pdf",
]


def _sobrenome_de(arquivo: str) -> str:
    """Deriva o sobrenome (1o token antes do '_') — igual ao verificar_autores."""
    return arquivo.split("_", 1)[0].replace("-", " ")


def _dica_topico(arquivo: str) -> str:
    """Para sobrenomes AMBIGUOS (mesmo autor, varios arquivos — ex.: Grewal
    Kalman + Grewal Power Electronics), extrai 2 palavras-chave do slug do
    titulo para desambiguar QUAL obra recuperar. Uma query so com o sobrenome
    e fragil nesse caso: o paper maior (Kalman, 1710 chunks) domina o menor
    (Power Electronics, 63 chunks). A dica de topico coincide com os gatilhos
    de TEXTBOOKS_PENALIZADOS, garantindo recuperacao deterministica."""
    base = arquivo.rsplit(".", 1)[0]
    partes = base.split("_")
    if len(partes) >= 3 and partes[-1].isdigit():
        meio = "_".join(partes[1:-1])
    else:
        meio = "_".join(partes[1:])
    palavras = [p for p in meio.replace("-", " ").split() if not p.isdigit()]
    return " ".join(palavras[:2]) if palavras else "o tema da dissertacao"


# ============================================================
# NOVO GERADOR — leitura de anexos (unit do leitor_anexos)
# ============================================================

def caso_anexo_texto(nome: str, arquivo: str, dados: bytes,
                     tokens: tuple[str, ...]) -> CasoTeste:
    def executar() -> tuple[bool, str]:
        r = ler_anexo(arquivo, dados)
        if r["tipo"] != "texto":
            return False, f"tipo={r['tipo']} (esperado texto); erro={r.get('erro')}"
        faltando = [t for t in tokens if t not in r["texto"]]
        ok = not faltando
        return ok, f"resumo={r['resumo']}; faltando={faltando or 'nenhum'}"

    return CasoTeste(nome, "leitura_anexos", arquivo, executar)


def caso_anexo_imagem(nome: str, arquivo: str, dados: bytes) -> CasoTeste:
    def executar() -> tuple[bool, str]:
        r = ler_anexo(arquivo, dados)
        ok = (
            r["tipo"] == "imagem"
            and bool(r["imagem_b64"])
            and r["mime"].startswith("image/")
        )
        return ok, (
            f"tipo={r['tipo']}; mime={r['mime']}; "
            f"b64_len={len(r['imagem_b64'])}; resumo={r['resumo']}"
        )

    return CasoTeste(nome, "leitura_anexos", arquivo, executar)


def caso_anexo_erro(nome: str, arquivo: str, dados: bytes,
                    erro_contem: str = "") -> CasoTeste:
    def executar() -> tuple[bool, str]:
        r = ler_anexo(arquivo, dados)
        if r["tipo"] != "erro":
            return False, f"tipo={r['tipo']} (esperado erro)"
        msg = f"{r.get('erro') or ''} {r.get('resumo') or ''}".lower()
        ok = (erro_contem.lower() in msg) if erro_contem else True
        return ok, f"erro={r.get('erro')}; resumo={r.get('resumo')}"

    return CasoTeste(nome, "leitura_anexos", arquivo, executar)


def caso_anexo_custom(nome: str, arquivo: str, dados: bytes, valida) -> CasoTeste:
    def executar() -> tuple[bool, str]:
        return valida(ler_anexo(arquivo, dados))

    return CasoTeste(nome, "leitura_anexos", arquivo, executar)


# ── Construtores de fixtures binarias em memoria ─────────────

def _fix_img(fmt: str, mode: str = "RGB", size: tuple[int, int] = (40, 40)) -> bytes:
    from PIL import Image
    if mode == "RGBA":
        cor = (200, 30, 30, 255)
    elif mode in ("P", "L"):
        cor = 100
    else:
        cor = (200, 30, 30)
    img = Image.new(mode, size, cor)
    buf = io.BytesIO()
    img.save(buf, format=fmt)
    return buf.getvalue()


def _fix_pdf_blank() -> bytes:
    from pypdf import PdfWriter
    w = PdfWriter()
    w.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    w.write(buf)
    return buf.getvalue()


def _fix_docx(paragrafos: tuple[str, ...], tabela: tuple[str, ...] = ()) -> bytes:
    import docx
    d = docx.Document()
    for p in paragrafos:
        d.add_paragraph(p)
    if tabela:
        t = d.add_table(rows=1, cols=len(tabela))
        for j, valor in enumerate(tabela):
            t.rows[0].cells[j].text = valor
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _fix_xlsx(abas: dict) -> bytes:
    from openpyxl import Workbook
    wb = Workbook()
    primeira = True
    for nome_aba, linhas in abas.items():
        ws = wb.active if primeira else wb.create_sheet(nome_aba)
        if primeira:
            ws.title = nome_aba
            primeira = False
        for linha in linhas:
            ws.append(linha)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _casos_leitura_anexos() -> list[CasoTeste]:
    """~65 unidades do leitor_anexos: texto, tabela, imagem, erro, borda."""
    casos: list[CasoTeste] = []

    # ── Texto / codigo / config (decode) ─────────────────────
    fix_texto: list[tuple[str, bytes, tuple[str, ...]]] = [
        ("relatorio.txt", b"Relatorio: tensao CA estavel, THD 2.1% no inversor.",
         ("inversor", "THD")),
        ("reuniao.txt", "Reuniao com a orientadora Fernanda sobre RCM e FMEA.".encode("utf-8"),
         ("orientadora", "RCM")),
        ("notas.md", "# Resultados\n\n- AUC=0.935\n- F1=0.98\n".encode("utf-8"),
         ("# Resultados", "AUC", "F1")),
        ("pipeline.md", "## Pipeline\nAutoencoder treinado com limiar p99.".encode("utf-8"),
         ("Autoencoder", "p99")),
        ("modelo.py", b"def detectar_anomalia(erro, limiar):\n    return erro > limiar\n",
         ("def detectar_anomalia", "limiar")),
        ("treino.py", b"import numpy as np\nthreshold = mu + 3*sigma\n",
         ("import numpy", "threshold")),
        ("app.js", b"function calcMTTF(falhas){ return 1/falhas; }",
         ("function calcMTTF",)),
        ("tipos.ts", b"const auc: number = 0.935;", ("const auc",)),
        ("Inversor.java", b"public class Inversor { int npr = 210; }",
         ("class Inversor", "210")),
        ("sinal.c", b"int main(){ float thd = 2.1; return 0; }", ("int main", "thd")),
        ("main.cpp", b"#include <iostream>\nint npr=210;", ("#include", "npr")),
        ("Prog.cs", b"class Prog { int npr = 210; }", ("class Prog",)),
        ("rul.go", b"package main\nfunc Weibull() float64 { return 0.0 }",
         ("package main", "Weibull")),
        ("vida.rs", b"fn rul() -> f64 { 1000.0 }", ("fn rul",)),
        ("calc.r", b"mttf <- function(x) mean(x)", ("mttf <- function",)),
        ("calc.rb", b"def npr; 210; end", ("def npr",)),
        ("modelo.kt", b"fun npr(): Int = 210", ("fun npr",)),
        ("v.swift", b"let auc = 0.935", ("let auc",)),
        ("m.scala", b"val rul = 1000", ("val rul",)),
        ("s.lua", b"local thd = 2.1", ("local thd",)),
        ("p.pl", b"my $npr = 210;", ("$npr",)),
        ("query.sql", b"SELECT npr FROM fmea WHERE componente='inversor';",
         ("SELECT npr", "inversor")),
        ("run.sh", b"#!/bin/bash\npython app.py --treinar", ("#!/bin/bash", "app.py")),
        ("deploy.ps1", b"$limiar = 2.91\nWrite-Output $limiar", ("$limiar",)),
        ("build.bat", b"@echo off\npython app.py", ("python app.py",)),
        ("idx.php", b"<?php echo 'inversor'; ?>", ("inversor",)),
        ("pagina.html", b"<html><body><h1>Relatorio PV</h1></body></html>",
         ("<h1>Relatorio PV</h1>",)),
        ("dados.xml", b"<fmea><npr>210</npr></fmea>", ("<npr>210</npr>",)),
        ("estilo.css", b".grafico { color: red; }", (".grafico",)),
        ("hdr.h", b"#define LIMIAR 2.91", ("#define LIMIAR",)),
        ("doc.tex", b"\\section{Resultados} AUC alto.", ("section{Resultados}",)),
        ("config.json", b'{"modelo": "autoencoder", "auc": 0.935}',
         ('"modelo"', "autoencoder")),
        ("lista.json", b'[1, 2, 3, "inversor"]', ("inversor",)),
        ("config.yaml", b"modelo: autoencoder\nauc: 0.935\n", ("modelo:", "autoencoder")),
        ("pipeline.yml", b"pipeline:\n  - features\n  - validacao\n",
         ("pipeline:", "features")),
        ("proj.toml", b'[projeto]\nnome = "mestrado"\n', ("[projeto]", "mestrado")),
        ("config.ini", b"[config]\nlimiar = 2.91\n", ("[config]", "limiar")),
        ("settings.cfg", b"debug=true\nprovedor=groq\n", ("provedor=groq",)),
        ("pipeline.log", b"2026-05-29 INFO pipeline iniciado\nERROR sensor CA\n",
         ("pipeline iniciado", "sensor CA")),
        ("leiame.rst", b"Titulo\n======\nConteudo RST do projeto.", ("Conteudo RST",)),
    ]
    for i, (arq, dados, tokens) in enumerate(fix_texto, 1):
        casos.append(caso_anexo_texto(f"anexo_txt_{i:02d}", arq, dados, tokens))

    # ── Tabelas: CSV / TSV ───────────────────────────────────
    fix_tab: list[tuple[str, bytes, tuple[str, ...]]] = [
        ("fmea.csv", b"componente,npr\ninversor,210\nsubsistema_ca,150\n",
         ("Colunas: componente, npr", "inversor")),
        ("metricas.csv", b"modelo,auc,f1\nautoencoder,1.0,0.98\n",
         ("Colunas: modelo, auc, f1", "autoencoder")),
        ("nums.csv", b"x,y\n1,2\n3,4\n", ("Colunas: x, y",)),
        ("cabecalho.csv", b"a,b,c\n", ("Colunas: a, b, c", "0 linhas")),
        ("tab.tsv", b"comp\tnpr\ninversor\t210\n", ("Colunas: comp, npr", "inversor")),
    ]
    for i, (arq, dados, tokens) in enumerate(fix_tab, 1):
        casos.append(caso_anexo_texto(f"anexo_tab_{i:02d}", arq, dados, tokens))

    # ── Excel (openpyxl) ─────────────────────────────────────
    casos.append(caso_anexo_texto(
        "anexo_xlsx_01", "planilha.xlsx",
        _fix_xlsx({"FMEA": [["componente", "npr"], ["inversor", 210], ["subsistema_ca", 150]]}),
        ("### Aba: FMEA", "componente", "inversor"),
    ))
    casos.append(caso_anexo_texto(
        "anexo_xlsx_02", "multi.xlsx",
        _fix_xlsx({
            "FMEA": [["componente", "npr"], ["inversor", 210]],
            "Metricas": [["metrica", "valor"], ["AUC", 0.935]],
        }),
        ("### Aba: FMEA", "### Aba: Metricas"),
    ))
    casos.append(caso_anexo_texto(
        "anexo_xlsx_03", "weibull.xlsx",
        _fix_xlsx({"RUL": [["t", "rul"], [100, 9000], [200, 8000]]}),
        ("### Aba: RUL", "Colunas: t, rul"),
    ))

    # ── Word (python-docx) ───────────────────────────────────
    casos.append(caso_anexo_texto(
        "anexo_docx_01", "relato.docx",
        _fix_docx(
            ("Resumo da dissertacao sobre inversores fotovoltaicos.",
             "Metodologia: FMEA, autoencoder e Weibull."),
            tabela=("componente", "NPR"),
        ),
        ("dissertacao", "autoencoder", "componente"),
    ))
    casos.append(caso_anexo_texto(
        "anexo_docx_02", "metodo.docx",
        _fix_docx(("O RCM orienta a manutencao centrada em confiabilidade.",
                   "O inversor e o componente mais critico (NPR=210).")),
        ("RCM", "NPR=210"),
    ))
    casos.append(caso_anexo_texto(
        "anexo_docx_03", "tabela.docx",
        _fix_docx(("Tabela de criticidade abaixo:",), tabela=("modo", "severidade")),
        ("criticidade", "modo", "severidade"),
    ))

    # ── Imagens (Pillow) ─────────────────────────────────────
    fix_img: list[tuple[str, bytes]] = [
        ("rgb.png", _fix_img("PNG", "RGB")),
        ("alfa.png", _fix_img("PNG", "RGBA")),
        ("grande.png", _fix_img("PNG", "RGB", size=(2000, 1500))),
        ("tiny.png", _fix_img("PNG", "RGB", size=(1, 1))),
        ("foto.jpg", _fix_img("JPEG", "RGB")),
        ("scan.jpeg", _fix_img("JPEG", "RGB", size=(300, 200))),
        ("anim.gif", _fix_img("GIF", "RGB")),
        ("bitmap.bmp", _fix_img("BMP", "RGB")),
        ("web.webp", _fix_img("WEBP", "RGB")),
    ]
    for i, (arq, dados) in enumerate(fix_img, 1):
        casos.append(caso_anexo_imagem(f"anexo_img_{i:02d}", arq, dados))

    # ── Erros e bordas ───────────────────────────────────────
    casos.append(caso_anexo_erro("anexo_err_01", "vazio.txt", b"", "vazio"))
    casos.append(caso_anexo_erro("anexo_err_02", "bin.xyz", b"\x00\x01\x02\x03" * 200,
                                 "nao suportado"))
    casos.append(caso_anexo_erro("anexo_err_03", "dados.dat", bytes(range(256)) * 4,
                                 "binario"))
    casos.append(caso_anexo_erro("anexo_err_04", "doc.pdf", _fix_pdf_blank(),
                                 "sem texto extraivel"))
    casos.append(caso_anexo_texto("anexo_err_05", "livre.qwerty",
                                  b"conteudo livre sem extensao conhecida",
                                  ("conteudo livre",)))
    casos.append(caso_anexo_texto("anexo_err_06", "anotacao.zzz",
                                  b"texto plano em extensao desconhecida",
                                  ("texto plano",)))

    def _valida_oversize(r: dict) -> tuple[bool, str]:
        ok = (
            r["tipo"] == "texto"
            and "truncado" in r["texto"]
            and len(r["texto"]) <= 12_200
        )
        return ok, f"tipo={r['tipo']}; len={len(r['texto'])}; tem_truncado={'truncado' in r['texto']}"

    casos.append(caso_anexo_custom("anexo_cap_01", "grande.txt",
                                   ("A" * 20_000).encode("utf-8"), _valida_oversize))
    casos.append(caso_anexo_custom("anexo_cap_02", "grande.md",
                                   ("# T\n" + ("linha de texto. " * 2_000)).encode("utf-8"),
                                   _valida_oversize))

    return casos


# ============================================================
# NOVO GERADOR — prompt com bloco de anexo
# ============================================================

_HDR_ANEXO = "ARQUIVOS ANEXADOS PELO PESQUISADOR (leia e use quando pertinente):"


def caso_prompt_com_anexo(nome: str,
                          pergunta: str,
                          anexos_texto: str,
                          consultar: bool,
                          deve_conter: tuple[str, ...],
                          nao_deve_conter: tuple[str, ...] = ()) -> CasoTeste:
    def executar() -> tuple[bool, str]:
        prompt = _montar_prompt(
            pergunta=pergunta,
            contexto="Memoria do projeto: resultados AUC=0,935.",
            historico_formatado="",
            orcamento=ORCAMENTO_TESTE,
            consultar_literatura=consultar,
            anexos_texto=anexos_texto,
        )
        faltando = [t for t in deve_conter if t not in prompt]
        indevidos = [t for t in nao_deve_conter if t in prompt]
        ok = not faltando and not indevidos
        return ok, f"faltando={faltando or 'nenhum'}; indevidos={indevidos or 'nenhum'}"

    return CasoTeste(nome, "prompt_com_anexo", pergunta, executar)


def _casos_prompt_anexo() -> list[CasoTeste]:
    """~18 casos: bloco de anexo presente quando ha anexo, ausente quando nao."""
    casos: list[CasoTeste] = []

    # Texto de anexo real, montado pelo pipeline (ler_anexos -> bloco).
    bloco_csv = montar_bloco_texto_anexos(
        ler_anexos([("fmea.csv", b"componente,npr\ninversor,210\n")])
    )
    bloco_txt = montar_bloco_texto_anexos(
        ler_anexos([("notas.txt", b"Reuniao sobre RCM e FMEA no inversor.")])
    )
    bloco_img_groq = montar_bloco_texto_anexos(
        ler_anexos([("grafico.png", _fix_img("PNG", "RGB"))]), suporta_imagem=False
    )

    com_anexo = [
        ("anexo_prompt_01", "O que tem nesse arquivo?",
         "[Anexo: dados.csv]\ncomponente,npr\ninversor,210",
         False, (_HDR_ANEXO, "componente,npr", "O pesquisador ANEXOU arquivos")),
        ("anexo_prompt_02", "Resuma o anexo.",
         "[Anexo: notas.txt]\nReuniao sobre RCM e FMEA.",
         False, (_HDR_ANEXO, "Reuniao sobre RCM", "Priorize esse conteudo")),
        ("anexo_prompt_03", "Analise os dados anexados.",
         bloco_csv, False, (_HDR_ANEXO, "componente", "inversor")),
        ("anexo_prompt_04", "Leia o documento.",
         bloco_txt, False, (_HDR_ANEXO, "RCM", "FMEA")),
        ("anexo_prompt_05", "Descreva a imagem anexada.",
         bloco_img_groq, False, (_HDR_ANEXO, "Gemini")),
        ("anexo_prompt_06", "Com base na literatura e no anexo, comente.",
         "[Anexo: tabela.csv]\nmodo,severidade\ncurto,5",
         True, (_HDR_ANEXO, "modo,severidade", "LITERATURA E MEMORIA")),
        ("anexo_prompt_07", "Interprete o arquivo.",
         "[Anexo: log.txt]\nERROR sensor CA fora da faixa",
         False, (_HDR_ANEXO, "sensor CA")),
        ("anexo_prompt_08", "Explique este codigo.",
         "[Anexo: modelo.py]\ndef detectar(): return True",
         False, (_HDR_ANEXO, "def detectar")),
        ("anexo_prompt_09", "Resuma a planilha.",
         "[Anexo: dados.xlsx]\n### Aba: FMEA\ncomponente | npr",
         False, (_HDR_ANEXO, "Aba: FMEA")),
        ("anexo_prompt_10", "O que diz o anexo sobre Weibull?",
         "[Anexo: rul.md]\nWeibull estima a vida util remanescente.",
         False, (_HDR_ANEXO, "Weibull")),
        ("anexo_prompt_11", "Analise com base no anexo e cite autores.",
         "[Anexo: resumo.txt]\nFMEA aplicado ao inversor.",
         True, (_HDR_ANEXO, "FMEA aplicado", "A pergunta pediu literatura/fontes")),
        ("anexo_prompt_12", "Veja o arquivo.",
         "[Anexo: config.json]\n{\"limiar\": 2.91}",
         False, (_HDR_ANEXO, "limiar")),
    ]
    for nome, pergunta, anexo, consultar, deve in com_anexo:
        casos.append(caso_prompt_com_anexo(nome, pergunta, anexo, consultar, deve))

    sem_anexo = [
        ("anexo_prompt_13", "Explique FMEA.", "", False),
        ("anexo_prompt_14", "Cite artigos sobre RCM.", "", True),
        ("anexo_prompt_15", "Interprete os resultados.", "", False),
        ("anexo_prompt_16", "Resuma o pipeline.", "   ", False),
        ("anexo_prompt_17", "Com base na literatura, fale de Weibull.", "", True),
        ("anexo_prompt_18", "Qual o proximo passo?", "\n\n", False),
    ]
    for nome, pergunta, anexo, consultar in sem_anexo:
        casos.append(caso_prompt_com_anexo(
            nome, pergunta, anexo, consultar, (), (_HDR_ANEXO, "ARQUIVOS ANEXADOS")
        ))

    return casos


def _casos_catalogo() -> list[CasoTeste]:
    """
    Catalogo da literatura (inventario completo). Blinda contra o bug de o
    agente listar so ~8 documentos e ALUCINAR o resto: pedir o inventario tem
    de rotear para a ferramenta deterministica `listar_base_bibliografica`,
    que le os metadados do ChromaDB (sem RAG, sem LLM) e nunca trunca/inventa.
    """
    casos: list[CasoTeste] = []

    # 1) Roteamento: pedidos de INVENTARIO -> listar_base_bibliografica
    frases_catalogo = [
        ("cat_route_01", "liste todas as referencias"),
        ("cat_route_02", "liste todas as suas referencias bibliograficas"),
        ("cat_route_03", "o que voce tem indexado?"),
        ("cat_route_04", "o que esta indexado na base?"),
        ("cat_route_05", "mostre a base bibliografica completa"),
        ("cat_route_06", "quantos artigos voce tem na base?"),
        ("cat_route_07", "quais documentos voce tem?"),
        ("cat_route_08", "mostre todas as 39 referencias"),
        ("cat_route_09", "liste a literatura indexada"),
        ("cat_route_10", "todas as fontes que voce possui"),
        ("cat_route_11", "me da o catalogo da base"),
        ("cat_route_12", "liste todos os artigos"),
        ("cat_route_13", "qual a bibliografia completa do projeto?"),
        ("cat_route_14", "quais referencias voce tem?"),
        ("cat_route_15", "me mostre a base de conhecimento inteira"),
        ("cat_route_16", "quantas obras existem na base?"),
    ]
    for nome, pergunta in frases_catalogo:
        casos.append(caso_ferramenta(nome, pergunta, "listar_base_bibliografica", True))

    # 2) Anti-regressao: buscas TEMATICAS nao podem virar catalogo (seguem RAG)
    frases_rag = [
        ("cat_nao_01", "cite artigos sobre deteccao de anomalias em inversores"),
        ("cat_nao_02", "o que a literatura diz sobre falhas no lado CA?"),
        ("cat_nao_03", "quais autores tratam de manutencao preditiva?"),
        ("cat_nao_04", "quais artigos voce tem sobre anomalias?"),
        ("cat_nao_05", "segundo a literatura, descreva o uso de Weibull"),
        ("cat_nao_06", "me fale sobre o artigo do Stender"),
        ("cat_nao_07", "referencias sobre RUL em eletronica de potencia"),
    ]
    for nome, pergunta in frases_rag:
        casos.append(caso_ferramenta(nome, pergunta, None, usar_esperado=False))

    # 3) Completude: catalogo_literatura lista TODOS os docs distintos da
    #    colecao, sem mojibake e sem extras (contagem == docs distintos reais).
    def executar_completo() -> tuple[bool, str]:
        _, colecao = _rag_cache()
        metas = (colecao.get(include=["metadatas"]).get("metadatas", []) or [])
        distintos = {m.get("arquivo") for m in metas if m.get("arquivo")}
        texto = catalogo_literatura(colecao)
        itens = [l for l in texto.splitlines() if l.startswith("- **")]
        sem_mojibake = "�" not in texto
        contagem_ok = len(itens) == len(distintos)
        total_ok = f"{len(distintos)} documentos indexados" in texto
        ok = bool(distintos) and contagem_ok and sem_mojibake and total_ok
        return ok, (
            f"distintos_colecao={len(distintos)}; itens_listados={len(itens)}; "
            f"sem_mojibake={sem_mojibake}; total_ok={total_ok}"
        )

    casos.append(CasoTeste(
        "cat_completude_contagem", "catalogo_literatura",
        "catalogo lista todos os docs distintos (sem truncar/inventar)",
        executar_completo,
    ))

    # 4) Despacho da ferramenta: resposta pronta, ok, e a lista chega completa.
    def executar_tool() -> tuple[bool, str]:
        _, colecao = _rag_cache()
        metas = (colecao.get(include=["metadatas"]).get("metadatas", []) or [])
        distintos = {m.get("arquivo") for m in metas if m.get("arquivo")}
        res = executar_ferramenta(
            "listar_base_bibliografica", pergunta="liste todas as referencias"
        )
        itens = [l for l in (res.get("mensagem") or "").splitlines()
                 if l.startswith("- **")]
        ok = (
            bool(res.get("ok"))
            and bool(res.get("resposta_pronta"))
            and len(itens) == len(distintos)
            and len(itens) >= 1
        )
        return ok, (
            f"ok={res.get('ok')}; resposta_pronta={res.get('resposta_pronta')}; "
            f"itens={len(itens)}; distintos={len(distintos)}"
        )

    casos.append(CasoTeste(
        "cat_tool_resposta_pronta", "catalogo_literatura",
        "ferramenta listar_base_bibliografica entrega o catalogo completo",
        executar_tool,
    ))

    return casos


def _casos_experimentos() -> list[CasoTeste]:
    """
    Experimentos de ML por artigo: roteamento ('rode o experimento do X',
    'quais experimentos existem') + integridade do registry (6 artigos, 5
    executáveis, Stender sem runner) + disponibilidade dos modelos. NÃO roda
    os experimentos pesados aqui — só valida estrutura e roteamento.
    """
    casos: list[CasoTeste] = []

    # "compare os experimentos" CONSULTA resultados salvos (roteamento
    # por intencao da sessao web); re-rodar exige "rode/execute".
    casos.append(caso_ferramenta(
        "exp_route_03", "compare os experimentos de anomalia",
        "consultar_resultados", True))

    # 1) Roteamento: RODAR experimento
    rodar = [
        ("exp_route_01", "rode o experimento do ghoneim"),
        ("exp_route_02", "teste os modelos do sharma"),
        ("exp_route_04", "rode todos os experimentos por artigo"),
        ("exp_route_05", "rode o experimento do francisti"),
        ("exp_route_06", "execute o experimento do ibrahim"),
        ("exp_route_07", "rode o experimento do ahirwar"),
    ]
    for nome, p in rodar:
        casos.append(caso_ferramenta(nome, p, "rodar_experimento_artigo", True))

    # 2) Roteamento: LISTAR experimentos
    listar = [
        ("exp_lista_01", "quais experimentos por artigo existem?"),
        ("exp_lista_02", "liste os experimentos disponiveis"),
        ("exp_lista_03", "mostre os experimentos por artigo"),
    ]
    for nome, p in listar:
        casos.append(caso_ferramenta(nome, p, "listar_experimentos_artigos", True))

    # 3) Anti-regressão: NÃO confundir com literatura/pipeline/RAG
    casos.append(caso_ferramenta(
        "exp_nao_01", "o que o ghoneim diz sobre deteccao de falhas?", None, False))
    casos.append(caso_ferramenta(
        "exp_nao_02", "rode o pipeline completo", "rodar_pipeline_completo", True))
    casos.append(caso_ferramenta(
        "exp_nao_03", "liste todas as referencias", "listar_base_bibliografica", True))

    # 4) Integridade do registry
    def executar_registry() -> tuple[bool, str]:
        from src.ml.experimentos_artigos import REGISTRO, listar_experimentos
        exps = listar_experimentos()
        execs = [e for e in exps if e.runner]
        stender = REGISTRO.get("stender")
        ok = (
            len(exps) == 6
            and len(execs) == 5
            and stender is not None
            and not stender.runner
            and all(e.modelos for e in execs)
        )
        return ok, (f"experimentos={len(exps)}; executaveis={len(execs)}; "
                    f"stender_sem_runner={stender is not None and not stender.runner}")

    casos.append(CasoTeste(
        "exp_registry_integro", "experimentos_artigos",
        "registry tem 6 artigos, 5 executáveis, Stender é cartão de dataset",
        executar_registry,
    ))

    # 5) Disponibilidade dos modelos (degradação honesta e coerente)
    def executar_disponibilidade() -> tuple[bool, str]:
        from src.ml.experimentos_artigos import lib_disponivel, listar_experimentos
        exps = listar_experimentos()
        total = sum(len(e.modelos) for e in exps)
        disp = sum(len(e.modelos_disponiveis()) for e in exps)
        # cada modelo indisponível deve declarar 'requer' uma lib real
        coerente = all(
            (m.requer is not None and not lib_disponivel(m.requer))
            for e in exps for m in e.modelos_indisponiveis()
        )
        ok = total >= 1 and disp >= 1 and coerente
        return ok, f"modelos_total={total}; disponiveis={disp}; degradacao_coerente={coerente}"

    casos.append(CasoTeste(
        "exp_disponibilidade", "experimentos_artigos",
        "modelos disponíveis e degradação coerente (indisponível => lib ausente)",
        executar_disponibilidade,
    ))

    return casos


def montar_casos() -> list[CasoTeste]:
    casos: list[CasoTeste] = []

    sem_literatura = [
        "Fale sobre FMEA.",
        "Explique RCM para eu aplicar no projeto.",
        "Interprete AUC=0,935 e F1=0,98.",
        "Qual o proximo passo da metodologia?",
        "Resuma o que fizemos no pipeline.",
        "Como eu explico Weibull na defesa?",
        "Me ajude a estruturar a secao de resultados.",
        "Quero estudar FMEA hoje.",
        "Analise a tabela de resultados do ML.",
        "Explique o papel do autoencoder.",
        "O que significa RUL?",
        "Sugira melhorias no texto da dissertacao.",
        "Compare validacao e treinamento.",
        "Como conecto FMECA com ML?",
        "Explique desbalanceamento no lado CA.",
        "Qual seria um bom paragrafo de conclusao?",
        "Revise minha ideia de experimento.",
        "O que o meu modelo aprendeu?",
        "Fale de falhas em inversores.",
        "Me ajude a planejar a proxima semana.",
    ]
    for i, pergunta in enumerate(sem_literatura + _gerar(_SEM_TPL, 7), 1):
        casos.append(caso_literatura(f"sem_literatura_{i:03d}", pergunta, False))

    com_literatura = [
        "Segundo a literatura, fale sobre FMEA.",
        "Com base na literatura, explique RCM.",
        "Consulte a literatura sobre Weibull em inversores PV.",
        "Fale sobre FMEA com referencias.",
        "Quais artigos falam de falhas em inversores?",
        "Cite autores sobre manutencao centrada em confiabilidade.",
        "Use fontes para explicar autoencoders.",
        "Liste referencias sobre diagnostico de falhas.",
        "O que a bibliografia diz sobre RUL?",
        "Existe paper sobre falhas CA em inversores?",
        "Faca uma revisao bibliografica de FMECA.",
        "Levante o estado da arte de manutencao preditiva PV.",
        "Use os documentos indexados para falar de IGBT.",
        "Na base de conhecimento, o que aparece sobre Weibull?",
        "Segundo os autores, como justificar FMEA?",
        "De acordo com os autores, qual a vantagem do RCM?",
        "Quero fontes sobre degradacao LCL.",
        "Procure referencias para digital twin.",
        "Faca um survey de ML preditivo.",
        "Traga um review de falhas fotovoltaicas.",
    ]
    for i, pergunta in enumerate(com_literatura + _gerar(_COM_TPL, 7), 1):
        casos.append(caso_literatura(f"com_literatura_{i:03d}", pergunta, True))

    ferramentas = [
        ("ferramenta_01", "Mostre os resultados do pipeline.", "consultar_resultados", True),
        ("ferramenta_02", "Quais metricas AUC e F1 foram obtidas?", "consultar_resultados", True),
        ("ferramenta_03", "Mostre os graficos da validacao.", "consultar_resultados", True),
        ("ferramenta_04", "Qual o status do pipeline?", "consultar_status_pipeline", True),
        ("ferramenta_05", "O que esta pendente no pipeline?", "consultar_status_pipeline", True),
        ("ferramenta_06", "Apague os resultados de ML para eu recalcular.", "limpar_resultados_ml", True),
        ("ferramenta_07", "Limpe os artefatos do autoencoder.", "limpar_resultados_ml", True),
        ("ferramenta_08", "Rode o pipeline completo.", "rodar_pipeline_completo", True),
        ("ferramenta_09", "Recalcule tudo do zero.", "rodar_pipeline_completo", True),
        ("ferramenta_10", "Rode as features CA.", "rodar_features_ca", True),
        ("ferramenta_11", "Treine o autoencoder.", "rodar_autoencoder", True),
        ("ferramenta_12", "Injete falhas sinteticas.", "rodar_injecao_falhas", True),
        ("ferramenta_13", "Valide com AUC e F1.", "rodar_validacao", True),
        ("ferramenta_14", "Calcule o Weibull.", "rodar_weibull", True),
        ("ferramenta_15", "Estime o RUL.", "rodar_weibull", True),
        ("ferramenta_16", "Busque na web a norma IEC 61724.", "buscar_web", True),
        ("ferramenta_17", "Pesquise na internet uma definicao oficial de FMEA.", "buscar_web", True),
        ("ferramenta_18", "Gere os resultados novamente.", "rodar_pipeline_completo", True),
        ("ferramenta_19", "Fale sobre FMEA.", None, False),
        ("ferramenta_20", "Explique o metodo de pesquisa.", None, False),
    ]
    # Extras com a ferramenta REAL capturada via decidir_acao(q, None) — nao
    # adivinhada. Captura sutilezas como "Calcule o MTTF" -> consultar_resultados
    # (porque "calcular" nao e substring de "calcule").
    ferramentas_extra = [
        ("ferramenta_21", "Qual o status atual do pipeline?", "consultar_status_pipeline", True),
        ("ferramenta_22", "O que ainda falta rodar?", "consultar_status_pipeline", True),
        ("ferramenta_23", "Mostre as curvas ROC.", "consultar_resultados", True),
        ("ferramenta_24", "Treine o detector de anomalias.", "rodar_autoencoder", True),
        ("ferramenta_25", "Calcule o MTTF e o B10.", "rodar_weibull", True),  # verbo de acao -> roda Weibull (roteamento por intencao)
        ("ferramenta_26", "Apague tudo e rode o pipeline de novo.", "limpar_resultados_ml", True),
        ("ferramenta_27", "Pesquise na internet a norma IEC 61724.", "buscar_web", True),
        ("ferramenta_28", "Mostre a matriz de confusao.", "consultar_resultados", True),
        ("ferramenta_29", "Rode a injecao de falhas sinteticas.", "rodar_injecao_falhas", True),
        ("ferramenta_30", "Recalcule tudo do zero.", "rodar_pipeline_completo", True),
        ("ferramenta_31", "Gere os graficos de validacao.", "rodar_validacao", True),
        ("ferramenta_32", "Estime a vida util remanescente (RUL).", "rodar_weibull", True),
        ("ferramenta_33", "Fale sobre a filosofia da manutencao.", None, False),
        ("ferramenta_34", "Explique o conceito de defesa em profundidade.", None, False),
    ]
    for nome, pergunta, ferramenta, usar in ferramentas + ferramentas_extra:
        casos.append(caso_ferramenta(nome, pergunta, ferramenta, usar))

    interacoes = [
        "Oi",
        "Opa",
        "Bom dia",
        "Boa tarde",
        "Boa noite",
        "valeu",
        "obrigado",
        "perfeito",
        "show",
        "beleza",
        "kkk",
        "entendi",
        "fechou",
        "tchau",
        "ate mais",
    ]
    interacoes_extra = [
        "ola", "salve", "eai", "hey", "alo", "fala", "obrigada", "vlw",
        "falou", "legal", "massa", "top", "ok", "certo", "combinado",
    ]
    for i, pergunta in enumerate(interacoes + interacoes_extra, 1):
        casos.append(caso_interacao(f"interacao_{i:02d}", pergunta))

    referencias = [
        ({}, ""),
        ({"a": "Torres (2024)"}, "- Torres (2024)"),
        ({"a": "Torres (2024)", "b": "Torres (2024)"}, "- Torres (2024)"),
        ({"a": "Torres (2024)", "b": "NASA (2008)"}, "- Torres (2024)\n- NASA (2008)"),
        (["A", "B", "A"], "- A\n- B"),
        ((None, "B"), "- B"),
        ({"a": "", "b": "C"}, "- C"),
        ({"a": "  D  "}, "- D"),
        (set(), ""),
        (["Fonte 1"], "- Fonte 1"),
        (["Fonte 1", "Fonte 2"], "- Fonte 1\n- Fonte 2"),
        ({"x": "Artigo X", "y": None, "z": "Artigo Z"}, "- Artigo X\n- Artigo Z"),
        (("Livro A", "Livro A", "Livro B"), "- Livro A\n- Livro B"),
        ({"k": 123}, "- 123"),
        ([False, "Fonte valida"], "- Fonte valida"),
    ]
    # Extras verificados contra a logica real (falsy 0/False/None/"" sao
    # ignorados; str().strip(); dedup preservando ordem).
    referencias_extra = [
        ({"a": "Smith (1999)"}, "- Smith (1999)"),
        (["Joshi (1996)", "Joshi (1996)"], "- Joshi (1996)"),
        ({"a": "Karim (2025)", "b": "Patil (2024)"}, "- Karim (2025)\n- Patil (2024)"),
        (("Voss (2009)",), "- Voss (2009)"),
        ({"a": None, "b": "Risi (2023)"}, "- Risi (2023)"),
        (["", "", "Sharma (2026)"], "- Sharma (2026)"),
        ({"a": "  Diniz (2021)  "}, "- Diniz (2021)"),
        ([0, "Pahwa (2017)"], "- Pahwa (2017)"),
        ({"a": "Moura (2019)", "b": "Moura (2019)", "c": "Silva (2008)"},
         "- Moura (2019)\n- Silva (2008)"),
        (("Frontin (2013)", None, "Frontin (2013)"), "- Frontin (2013)"),
        ({"k": 2024}, "- 2024"),
        ([0, 0, 0], ""),
        ({"a": 0, "b": False, "c": None}, ""),
        (["Xavier (2005)", "", "Dhople (2012)", "Xavier (2005)"],
         "- Xavier (2005)\n- Dhople (2012)"),
        ({"x": "Oppenheim (2014)"}, "- Oppenheim (2014)"),
        (("Gonzalez (2008)", "Tekalp (2015)"), "- Gonzalez (2008)\n- Tekalp (2015)"),
        ({"a": "  ", "b": "Stewart (2013)"}, "- Stewart (2013)"),
        (["Sakurada (1998)"], "- Sakurada (1998)"),
        ({"a": "Cristaldi (2017)", "b": "Cristaldi (2017)"}, "- Cristaldi (2017)"),
        ([42, "Item"], "- 42\n- Item"),
    ]
    for i, (entrada, esperado) in enumerate(referencias + referencias_extra, 1):
        casos.append(caso_referencias(f"referencia_{i:02d}", entrada, esperado))

    prompts = [
        (
            "prompt_01",
            "Fale sobre FMEA.",
            False,
            ("CONTEXTO RECUPERADO DA MEMORIA DO PROJETO", "NAO pediu literatura/fontes"),
            ("CONTEXTO RECUPERADO DA LITERATURA E MEMORIA",),
        ),
        (
            "prompt_02",
            "Com base na literatura, fale sobre FMEA.",
            True,
            ("CONTEXTO RECUPERADO DA LITERATURA E MEMORIA", "A pergunta pediu literatura/fontes"),
            ("CONTEXTO RECUPERADO DA MEMORIA DO PROJETO",),
        ),
        (
            "prompt_03",
            "Interprete meus resultados.",
            False,
            ("Use apenas conhecimento do projeto",),
            ("As referencias externas serao exibidas",),
        ),
        (
            "prompt_04",
            "Cite artigos sobre RCM.",
            True,
            ("cite autor/ano", "lista de fontes consultadas"),
            (),
        ),
        (
            "prompt_05",
            "Explique autoencoder.",
            False,
            ("sem mencionar literatura",),
            (),
        ),
        (
            "prompt_06",
            "Quais fontes explicam Weibull?",
            True,
            ("evidencias recuperadas",),
            (),
        ),
        (
            "prompt_07",
            "Ajude no texto da defesa.",
            False,
            ("MEMORIA DO PROJETO",),
            ("LITERATURA E MEMORIA",),
        ),
        (
            "prompt_08",
            "Segundo os autores, explique FMECA.",
            True,
            ("LITERATURA E MEMORIA",),
            (),
        ),
        (
            "prompt_09",
            "O que significa RUL?",
            False,
            ("raciocinio tecnico geral",),
            (),
        ),
        (
            "prompt_10",
            "Faca uma revisao bibliografica sobre falhas CA.",
            True,
            ("NUNCA escreva uma secao final",),
            (),
        ),
    ]
    prompts_extra = [
        ("prompt_11", "Explique o autoencoder do projeto.", False,
         ("CONTEXTO RECUPERADO DA MEMORIA DO PROJETO",),
         ("CONTEXTO RECUPERADO DA LITERATURA E MEMORIA",)),
        ("prompt_12", "Resuma os resultados de validacao.", False,
         ("MEMORIA DO PROJETO",), ("LITERATURA E MEMORIA",)),
        ("prompt_13", "Como conecto FMECA e ML?", False,
         ("raciocinio tecnico geral",), ()),
        ("prompt_14", "Qual o proximo passo do pipeline?", False,
         ("sem mencionar literatura",), ()),
        ("prompt_15", "Interprete o limiar p99.", False,
         ("MEMORIA DO PROJETO",), ()),
        ("prompt_16", "Ajude-me a planejar a semana.", False,
         ("MEMORIA DO PROJETO",), ()),
        ("prompt_17", "Explique desbalanceamento no lado CA.", False,
         ("MEMORIA DO PROJETO",), ()),
        ("prompt_18", "Cite artigos sobre autoencoder.", True,
         ("LITERATURA E MEMORIA", "A pergunta pediu literatura/fontes"), ()),
        ("prompt_19", "Segundo a literatura, explique Weibull.", True,
         ("LITERATURA E MEMORIA",), ()),
        ("prompt_20", "Liste referencias sobre RCM.", True,
         ("cite autor/ano",), ()),
        ("prompt_21", "Com base na literatura, descreva FMECA.", True,
         ("LITERATURA E MEMORIA",), ()),
        ("prompt_22", "Faca uma revisao bibliografica de RUL.", True,
         ("NUNCA escreva uma secao final",), ()),
        ("prompt_23", "Quais papers tratam de isolation forest?", True,
         ("LITERATURA E MEMORIA",), ()),
        ("prompt_24", "O que a bibliografia diz sobre IGBT?", True,
         ("A pergunta pediu literatura/fontes",), ()),
    ]
    for nome, pergunta, consultar, deve, nao_deve in prompts + prompts_extra:
        casos.append(caso_prompt(nome, pergunta, consultar, deve, nao_deve))

    # ── Diversidade real de literatura (RAG end-to-end) ──────
    # Textbooks que nao devem aparecer em queries amplas: math/imagem/video/dsp.
    textbooks_off = (
        "stewart_calculo-volume-i_2013.pdf",
        "gonzalez_digital-image-processing_2008.pdf",
        "tekalp_digital-video-processing_2015.pdf",
    )
    casos_diversidade = [
        (
            "diversidade_01",
            "Faça uma literatura completa da dissertação. cite a literatura.",
            6,
            textbooks_off,
        ),
        (
            "diversidade_02",
            "Faça uma revisao bibliografica sobre FMEA em inversor fotovoltaico.",
            6,
            textbooks_off,
        ),
        (
            "diversidade_03",
            "cite a literatura sobre autoencoder em inversor PV",
            5,
            textbooks_off,
        ),
        (
            "diversidade_04",
            "Levante o estado da arte de manutencao preditiva em sistemas PV.",
            6,
            textbooks_off,
        ),
        (
            "diversidade_05",
            "Cite os artigos sobre Weibull e RUL no contexto fotovoltaico.",
            5,
            textbooks_off,
        ),
    ]
    for nome, pergunta, min_fontes, proibidos in casos_diversidade:
        casos.append(caso_diversidade_literatura(nome, pergunta, min_fontes, arquivos_proibidos=proibidos))

    # ── Strip de blocos de fontes hallucinados pelo LLM ──────
    casos_strip = [
        (
            "strip_01",
            "Resposta do agente.\n\n---\n📚 **Fontes:**\n- Torres (2024)\n- Stender (2020)\n",
            ("📚", "Torres (2024)", "Stender (2020)"),
            ("Resposta do agente.",),
        ),
        (
            "strip_02",
            "Texto principal com (Torres, 2024) citado.\n\n## Referências\n\n- Torres, R. (2024)\n",
            ("## Referências", "Torres, R. (2024)"),
            ("Texto principal", "(Torres, 2024)"),
        ),
        (
            "strip_03",
            "Conclusao do paragrafo.\n\n**Referências bibliográficas:**\n- A\n- B\n",
            ("**Referências bibliográficas:**", "- A", "- B"),
            ("Conclusao do paragrafo.",),
        ),
        (
            "strip_04",
            "Resposta limpa sem bloco de fontes anexado.",
            (),
            ("Resposta limpa sem bloco",),
        ),
        (
            "strip_05",
            "Analise tecnica.\n\n### Bibliografia\n- Artigo 1\n- Artigo 2\n",
            ("### Bibliografia", "Artigo 1", "Artigo 2"),
            ("Analise tecnica.",),
        ),
        (
            "strip_06",
            "Texto.\n\n---\n\n📚 **Fontes consultadas:**\n- X\n",
            ("📚", "Fontes consultadas", "- X"),
            ("Texto.",),
        ),
        (
            "strip_07",
            "Texto com (Autor, 2024) inline.\n\nREFERÊNCIAS:\n- Autor (2024)\n",
            ("REFERÊNCIAS:", "- Autor (2024)"),
            ("Texto com (Autor, 2024) inline.",),
        ),
        # ── Mais cabecalhos que DEVEM ser cortados (header + lista) ──
        (
            "strip_08",
            "Analise tecnica do resultado.\n\n## Bibliografia\n- A\n- B\n",
            ("## Bibliografia", "- A", "- B"),
            ("Analise tecnica do resultado.",),
        ),
        (
            "strip_09",
            "Texto final aqui.\n\nFontes:\n- X\n- Y\n",
            ("Fontes:", "- X", "- Y"),
            ("Texto final aqui.",),
        ),
        (
            "strip_10",
            "Corpo da resposta.\n\n### Referências bibliográficas\n* Item 1\n* Item 2\n",
            ("### Referências bibliográficas", "Item 1", "Item 2"),
            ("Corpo da resposta.",),
        ),
        (
            "strip_11",
            "Resposta concluida.\n\n📚 Fontes consultadas:\n1. Torres (2024)\n",
            ("📚", "Torres (2024)"),
            ("Resposta concluida.",),
        ),
        (
            "strip_12",
            "Discussao do paragrafo.\n\n**Referências:**\n- A\n",
            ("**Referências:**", "- A"),
            ("Discussao do paragrafo.",),
        ),
        (
            "strip_13",
            "Conclusao do texto.\n\nReferencias:\n- Item unico\n",
            ("Referencias:", "- Item unico"),
            ("Conclusao do texto.",),
        ),
        (
            "strip_14",
            "Paragrafo final.\n\nBibliografia:\n- Fonte 1\n- Fonte 2\n",
            ("Bibliografia:", "- Fonte 1"),
            ("Paragrafo final.",),
        ),
        (
            "strip_15",
            "Analise pronta.\n\n## Fontes\n- A\n",
            ("## Fontes", "- A"),
            ("Analise pronta.",),
        ),
        (
            "strip_16",
            "Texto base.\n\n**Bibliografia**\n- A\n- B\n",
            ("**Bibliografia**", "- A"),
            ("Texto base.",),
        ),
        (
            "strip_17",
            # Separador '---' antes do bloco deve ser engolido.
            "Texto principal.\n\n---\n## Referências\n- A\n",
            ("## Referências", "- A", "---"),
            ("Texto principal.",),
        ),
        (
            "strip_18",
            # Blocos consecutivos: todos somem a partir do primeiro.
            "Texto.\n\n## Referencias\n- A\n\n📚 Fontes:\n- B\n",
            ("## Referencias", "📚 Fontes", "- A", "- B"),
            ("Texto.",),
        ),
        (
            "strip_19",
            "Resposta tecnica completa.\n\nREFERÊNCIAS:\n- Autor (2024)\n- Outro (2023)\n",
            ("REFERÊNCIAS:", "Autor (2024)", "Outro (2023)"),
            ("Resposta tecnica completa.",),
        ),
        # ── Casos que NAO devem ser cortados (preservacao) ──
        (
            "strip_20",
            "Conteudo sem nenhum bloco de fontes anexado ao final.",
            (),
            ("Conteudo sem nenhum bloco",),
        ),
        (
            "strip_21",
            "O paper (Torres, 2024) discute FMEA inline no corpo. Fim do texto.",
            (),
            ("(Torres, 2024)", "Fim do texto."),
        ),
        (
            "strip_22",
            # 📚 + 'Fontes' em prosa NO MEIO, seguido de mais prosa: nao corta.
            "Intro do paragrafo aqui.\n\n📚 Fontes foram discutidas acima no texto.\n\nConclusao final do paragrafo segue normalmente.",
            (),
            ("Conclusao final do paragrafo",),
        ),
        (
            "strip_23",
            "O metodo usa varias fontes de dados industriais para treinar o modelo.",
            (),
            ("fontes de dados industriais",),
        ),
        (
            "strip_24",
            # 'Passos:' nao e cabecalho de fontes — lista numerada deve sobreviver.
            "Plano de execucao.\n\nPassos do pipeline:\n1. extrair features\n2. treinar\n",
            (),
            ("extrair features", "treinar"),
        ),
    ]
    for nome, entrada, deve_remover, deve_preservar in casos_strip:
        casos.append(caso_strip_fontes(nome, entrada, deve_remover, deve_preservar))

    # ── Casos adversariais de anti-ruido ─────────────────────
    # Queries que ANTIGAMENTE traziam textbooks irrelevantes; agora devem trazer
    # apenas papers do dominio.
    casos_adversarial = [
        (
            "adversarial_01",
            # 'calculo' no texto da query NAO pode promover Stewart Calculo.
            "Explique o calculo do limiar p99 do autoencoder.",
            3,
            ("stewart_calculo-volume-i_2013.pdf",),
        ),
        (
            "adversarial_02",
            # 'imagem' no contexto de termografia: ainda assim Gonzalez nao deve dominar.
            "Como sao analisadas as imagens termograficas de inversores PV?",
            3,
            ("tekalp_digital-video-processing_2015.pdf",),
        ),
        (
            "adversarial_03",
            # '+30 artigos' nao pode pegar tabela com numero 30 em livro irrelevante.
            "Tem +30 artigos sobre falhas em inversores indexados?",
            3,
            ("stewart_calculo-volume-i_2013.pdf", "gonzalez_digital-image-processing_2008.pdf"),
        ),
        (
            "adversarial_04",
            # Pergunta vaga sobre tabelas — antes virava ruido absoluto.
            "Quais tabelas a literatura traz sobre criticidade de inversor?",
            4,
            ("stewart_calculo-volume-i_2013.pdf",),
        ),
        (
            "adversarial_05",
            # 'metodologia' sozinha — pergunta ampla.
            "Descreva a metodologia da literatura sobre FMECA.",
            4,
            ("stewart_calculo-volume-i_2013.pdf", "tekalp_digital-video-processing_2015.pdf"),
        ),
    ]
    for nome, pergunta, min_fontes, proibidos in casos_adversarial:
        casos.append(caso_diversidade_literatura(
            f"adversarial_{nome.split('_')[1]}", pergunta, min_fontes, arquivos_proibidos=proibidos
        ))

    # ── Strip robusto: casos esquisitos ─────────────────────
    casos_strip_extra = [
        (
            "strip_extra_01",
            # Bloco no MEIO do texto NAO deve ser cortado.
            "Intro do paragrafo.\n\n📚 Fontes do paragrafo anterior estavam ok.\n\nContinuacao do texto.",
            (),
            ("Continuacao do texto.",),
        ),
        (
            "strip_extra_02",
            # Citacoes inline (Autor, ano) DEVEM ser preservadas.
            "O FMEA foi descrito por (Stamatis, 2003) e refinado por (Torres, 2024). Fim.",
            (),
            ("(Stamatis, 2003)", "(Torres, 2024)", "Fim."),
        ),
        (
            "strip_extra_03",
            # Multiplos blocos consecutivos: todos devem sumir.
            "Texto.\n\n## Referencias\n- A\n\n📚 Fontes:\n- B\n",
            ("## Referencias", "📚 Fontes", "- A", "- B"),
            ("Texto.",),
        ),
    ]
    for nome, entrada, deve_remover, deve_preservar in casos_strip_extra:
        casos.append(caso_strip_fontes(nome, entrada, deve_remover, deve_preservar))

    # ── Proveniencia topica: paper correto para query especifica ──
    # Estes testes blindam contra regressoes onde o RAG perde os papers
    # core de cada area da dissertacao.
    casos_proveniencia = [
        (
            "proveniencia_01",
            "explique autoencoder para deteccao de anomalias em inversor PV",
            (
                "francisti_predictive-modeling-and-anomaly-detection-in-solar-pv-invert_2025.pdf",
                "ahirwar_enhanced-anomaly-detection-in-solar-power-plants-using-hybri_2025.pdf",
                "ibrahim_machine-learning-schemes-for-anomaly-detection-in-solar-powe_2022.pdf",
            ),
            2,
        ),
        (
            "proveniencia_02",
            "Weibull confiabilidade RUL inversor fotovoltaico",
            (
                "karim_a-review-on-risk-and-reliability-analysis-in-photovoltaic-po_2025.pdf",
                "patil_a-reliability-and-risk-assessment-of-solar-photovoltaic-pane_2024.pdf",
                "shuttleworth_reliability-prediction-of-pv-inverters-based-on-mil-hdbk-217_2015.pdf",
                "dhople_estimation-of-photovoltaic-system-reliability-and-performanc_2012.pdf",
            ),
            2,
        ),
        (
            "proveniencia_03",
            "FMEA FMECA NPR criticidade inversor",
            (
                "torres_aplicacao-da-metodologia-reliability-centred-maintenance-a-s_2024.pdf",
                "sakurada_as-tecnicas-de-analise-do-modos-de-falhas-e-seus-efeitos-e-a_1998.pdf",
            ),
            2,
        ),
        (
            "proveniencia_04",
            "IGBT failure inverter dataset Paderborn",
            (
                "stender_data-set-description-three-phase-igbt-two-level-inverter-for_2020.pdf",
                "cristaldi_a-root-cause-analysis-and-a-risk-evaluation-of-pv-balance-of_2017.pdf",
            ),
            1,
        ),
        (
            "proveniencia_05",
            "Random Forest XGBoost fault detection PV inverter",
            (
                "ghoneim_fault-detection-algorithms-for-achieving-service-continuity-_2021.pdf",
                "narayanan_machine-learning-based-explainable-fault-detection-of-vacuum_2023.pdf",
                "ibrahim_machine-learning-schemes-for-anomaly-detection-in-solar-powe_2022.pdf",
            ),
            2,
        ),
        (
            "proveniencia_06",
            "RCM manutencao centrada em confiabilidade",
            (
                "torres_aplicacao-da-metodologia-reliability-centred-maintenance-a-s_2024.pdf",
                "administration_nasa-reliability-centered-maintenance-guide-for-facilities-a_2008.pdf",
                "muqauwim_analysis-of-optimal-maintenance-interval-on-id-fan-using-rel_2020.pdf",
            ),
            2,
        ),
        (
            "proveniencia_07",
            "isolation forest deteccao anomalias",
            (
                "sharma_a-self-tuning-reinforcement-learning-driven-isolation-forest_2026.pdf",
                "ibrahim_machine-learning-schemes-for-anomaly-detection-in-solar-powe_2022.pdf",
            ),
            1,
        ),
    ]
    for nome, pergunta, esperados, min_match in casos_proveniencia:
        casos.append(caso_proveniencia(nome, pergunta, esperados, min_match))

    # ── Trigger por autor + recuperacao do autor ─────────────
    # Cenario reportado pelo Rodolfo: "E o da NASA?" nao disparava RAG
    # e o LLM dizia "minha base nao tem NASA" — falsamente.
    def caso_autor_trigger_e_recupera(nome, pergunta, autor_arquivo_esperado):
        def executar():
            try:
                modelo, colecao = _rag_cache()
                # Trigger: deve_consultar_literatura precisa dizer True
                consultar = deve_consultar_literatura(pergunta, colecao)
                if not consultar:
                    return False, f"deve_consultar_literatura={consultar}; deveria=True"
                # Recuperacao: o autor citado precisa aparecer no top
                ctx, citacoes = buscar_contexto(
                    pergunta, modelo, colecao,
                    n_pool=120, n_resultados=16, n_resultados_revisao=28,
                    max_chunks_por_fonte=2, contexto_chars=14_000,
                    sessao_chars=1_500, consultar_literatura=True,
                )
                fontes = {str(k).split('|')[0] for k in citacoes}
                achou = autor_arquivo_esperado in fontes
                return achou, (
                    f"consultar=True; autor_recuperado={achou}; "
                    f"top={list(citacoes.keys())[:3] or 'nenhum'}"
                )
            except Exception as exc:
                return False, f"excecao={type(exc).__name__}: {exc}"

        return CasoTeste(nome, "autor_trigger_rag", pergunta, executar)

    casos_autor = [
        (
            "autor_01_nasa_cade",
            "E o da nasa? cadê?",
            "administration_nasa-reliability-centered-maintenance-guide-for-facilities-a_2008.pdf",
        ),
        (
            "autor_02_torres_cade",
            "Cadê o Torres?",
            "torres_aplicacao-da-metodologia-reliability-centred-maintenance-a-s_2024.pdf",
        ),
        (
            "autor_03_ahirwar_tem",
            "tem o Ahirwar na base?",
            "ahirwar_enhanced-anomaly-detection-in-solar-power-plants-using-hybri_2025.pdf",
        ),
        (
            "autor_04_stender_fala",
            "o que o Stender fala?",
            "stender_data-set-description-three-phase-igbt-two-level-inverter-for_2020.pdf",
        ),
        (
            "autor_05_lafraia_indexacao",
            "perdi a indexação do Lafraia?",
            "lafraia_manual-de-confiabilidade-mantenabilidade-e-disponibilidade_0000.pdf",
        ),
        (
            "autor_06_francisti_paper",
            "cite o paper do Francisti",
            "francisti_predictive-modeling-and-anomaly-detection-in-solar-pv-invert_2025.pdf",
        ),
        (
            "autor_07_karim_indexado",
            "tem alguma coisa de Karim indexada?",
            "karim_a-review-on-risk-and-reliability-analysis-in-photovoltaic-po_2025.pdf",
        ),
        (
            "autor_08_monteiro_que_diz",
            "o que o Monteiro diz sobre falhas em inversores?",
            "monteiro_identifying-critical-failures-in-pv-systems-based-on-pv-inve_2024.pdf",
        ),
        (
            "autor_09_ghoneim_fala",
            "Ghoneim fala de quê?",
            "ghoneim_fault-detection-algorithms-for-achieving-service-continuity-_2021.pdf",
        ),
        (
            "autor_10_paderborn_dataset",
            "tem o dataset Paderborn indexado?",
            "stender_data-set-description-three-phase-igbt-two-level-inverter-for_2020.pdf",
        ),
    ]
    for nome, pergunta, arq in casos_autor:
        casos.append(caso_autor_trigger_e_recupera(nome, pergunta, arq))

    # ── Cobertura SISTEMATICA dos 39 documentos ──────────────
    # Para cada arquivo indexado: dispara RAG pelo sobrenome e exige que o
    # proprio arquivo volte. Blinda "alem do NASA, os outros 38 aparecem?".
    from collections import Counter
    _freq_sobrenome = Counter(_sobrenome_de(a) for a in ARQUIVOS_INDEXADOS)
    for idx, arquivo in enumerate(ARQUIVOS_INDEXADOS, 1):
        sobrenome = _sobrenome_de(arquivo)
        if _freq_sobrenome[sobrenome] > 1:
            # Sobrenome compartilhado por >1 arquivo: desambigua pelo topico.
            pergunta = f"o que o {sobrenome} diz sobre {_dica_topico(arquivo)}?"
        else:
            pergunta = f"o que o {sobrenome} diz sobre o tema da dissertacao?"
        casos.append(caso_autor_trigger_e_recupera(
            f"autor_sys_{idx:02d}_{sobrenome.replace(' ', '_')}", pergunta, arquivo
        ))

    # ── Prompt anti-alucinacao: contem instrucao de nao afirmar ausencia
    casos.append(caso_prompt(
        "prompt_anti_alucinacao_01",
        "tem o Stender na base?",
        True,
        ("NUNCA afirme que um autor", "nao veio agora na minha busca"),
        (),
    ))
    casos.append(caso_prompt(
        "prompt_anti_alucinacao_02",
        "cadê o paper do Ahirwar?",
        True,
        ("NUNCA afirme",),
        (),
    ))
    casos.append(caso_prompt(
        "prompt_anti_alucinacao_03",
        "tem o Karim na base?",
        True,
        ("NUNCA afirme que um autor",),
        (),
    ))
    casos.append(caso_prompt(
        "prompt_anti_alucinacao_04",
        "cadê o Monteiro?",
        True,
        ("nao veio agora na minha busca",),
        (),
    ))
    casos.append(caso_prompt(
        "prompt_anti_alucinacao_05",
        "perdi o Ghoneim?",
        True,
        ("Nunca afirme ausencia total",),
        (),
    ))
    casos.append(caso_prompt(
        "prompt_anti_alucinacao_06",
        "tem algo do Stender indexado?",
        True,
        ("nao esta na base",),
        (),
    ))

    # ── Stress / borda ───────────────────────────────────────
    def caso_stress(nome, executor):
        return CasoTeste(nome, "stress_borda", "(stress test)", executor)

    def stress_pergunta_vazia():
        try:
            modelo, colecao = _rag_cache()
            ctx, cit = buscar_contexto(
                "", modelo, colecao,
                n_pool=60, n_resultados=8, max_chunks_por_fonte=2,
                contexto_chars=5000, sessao_chars=500,
                consultar_literatura=True,
            )
            # Pergunta vazia: nao explode mas tambem nao retorna ruido relevante
            return True, f"sem_excecao; n_fontes={len(cit)}"
        except Exception as exc:
            return False, f"excecao={type(exc).__name__}: {exc}"

    def stress_pergunta_longa():
        try:
            pergunta = ("FMEA " * 500) + "autoencoder"
            modelo, colecao = _rag_cache()
            ctx, cit = buscar_contexto(
                pergunta, modelo, colecao,
                n_pool=60, n_resultados=8, max_chunks_por_fonte=2,
                contexto_chars=5000, sessao_chars=500,
                consultar_literatura=True,
            )
            return len(cit) > 0, f"n_fontes={len(cit)}"
        except Exception as exc:
            return False, f"excecao={type(exc).__name__}: {exc}"

    def stress_strip_empty():
        try:
            r1 = remover_bloco_fontes_llm("")
            r2 = remover_bloco_fontes_llm(None)
            r3 = remover_bloco_fontes_llm("   ")
            ok = r1 == "" and r2 is None and r3.strip() == ""
            return ok, f"r1={r1!r}; r2={r2!r}; r3={r3!r}"
        except Exception as exc:
            return False, f"excecao={type(exc).__name__}: {exc}"

    def stress_strip_apenas_bloco():
        """Resposta inteira é só uma 'Referencias:' — deve virar string vazia."""
        try:
            entrada = "## Referências\n- A\n- B\n"
            saida = remover_bloco_fontes_llm(entrada)
            ok = saida.strip() == "" or "##" not in saida
            return ok, f"saida={saida!r}"
        except Exception as exc:
            return False, f"excecao={type(exc).__name__}: {exc}"

    def stress_strip_caracteres_especiais():
        entrada = "Texto com ç ã é à ô — emoji 🔬⚡📊.\n\n📚 **Fontes:**\n- Fonte (2024)\n"
        try:
            saida = remover_bloco_fontes_llm(entrada)
            ok = "📚" not in saida and "ç ã é" in saida and "🔬" in saida
            return ok, f"acentos_preservados={('ç ã é' in saida)}; bloco_removido={'📚' not in saida}"
        except Exception as exc:
            return False, f"excecao={type(exc).__name__}: {exc}"

    def stress_strip_idempotente():
        """Aplicar strip duas vezes da mesmo resultado da primeira."""
        entrada = "Texto.\n\n📚 **Fontes:**\n- A\n"
        try:
            uma = remover_bloco_fontes_llm(entrada)
            duas = remover_bloco_fontes_llm(uma)
            ok = uma == duas
            return ok, f"idempotente={ok}"
        except Exception as exc:
            return False, f"excecao={type(exc).__name__}: {exc}"

    def stress_orcamento_groq_revisao():
        """Para query de revisao em provedor Groq, deve trazer >=4 fontes apesar do orcamento menor."""
        try:
            modelo, colecao = _rag_cache()
            ctx, cit = buscar_contexto(
                "revisao bibliografica de FMEA em PV",
                modelo, colecao,
                n_pool=60, n_resultados=10, n_resultados_revisao=16,
                max_chunks_por_fonte=2, contexto_chars=7000, sessao_chars=800,
                consultar_literatura=True,
            )
            ok = len(cit) >= 4
            return ok, f"n_fontes={len(cit)}"
        except Exception as exc:
            return False, f"excecao={type(exc).__name__}: {exc}"

    def stress_sem_literatura_nao_consulta():
        """consultar_literatura=False deve retornar contexto SEM bloco da literatura."""
        try:
            modelo, colecao = _rag_cache()
            ctx, cit = buscar_contexto(
                "Fale sobre FMEA.", modelo, colecao,
                n_pool=30, n_resultados=8, max_chunks_por_fonte=2,
                contexto_chars=5000, sessao_chars=500,
                consultar_literatura=False,
            )
            ok = "DA LITERATURA" not in ctx and len(cit) == 0
            return ok, f"sem_literatura={ok}"
        except Exception as exc:
            return False, f"excecao={type(exc).__name__}: {exc}"

    stresses = [
        ("stress_01_pergunta_vazia", stress_pergunta_vazia),
        ("stress_02_pergunta_longa", stress_pergunta_longa),
        ("stress_03_strip_empty", stress_strip_empty),
        ("stress_04_strip_apenas_bloco", stress_strip_apenas_bloco),
        ("stress_05_strip_caracteres_especiais", stress_strip_caracteres_especiais),
        ("stress_06_strip_idempotente", stress_strip_idempotente),
        ("stress_07_orcamento_groq_revisao", stress_orcamento_groq_revisao),
        ("stress_08_sem_literatura_nao_consulta", stress_sem_literatura_nao_consulta),
    ]
    for nome, fn in stresses:
        casos.append(caso_stress(nome, fn))

    # ── Contexto diversificado (o LLM recebe varios [Fonte:]) ──
    casos_contexto = [
        ("contexto_01", "Faça uma literatura completa da dissertação.", 6),
        ("contexto_02", "cite a literatura sobre autoencoder em inversor PV", 5),
        ("contexto_03", "revisao bibliografica sobre Weibull em PV", 5),
        ("contexto_04", "Com base na literatura, explique FMEA.", 4),
        ("contexto_05", "Levante o estado da arte de manutencao preditiva.", 5),
    ]
    for nome, pergunta, min_fontes in casos_contexto:
        casos.append(caso_contexto_diverso(nome, pergunta, min_fontes))

    # ── Detector de query de revisao + expansao de topicos ──
    casos_revisao = [
        ("revisao_01", "Faca uma revisao bibliografica de FMEA.", True, ("FMEA",)),
        ("revisao_02", "Cite a literatura completa da dissertação.", True, ("autoencoder",)),
        ("revisao_03", "Levante o estado da arte de ML preditivo em PV.", True, ("Weibull",)),
        ("revisao_04", "Faca um survey de manutencao preditiva.", True, ("RCM",)),
        ("revisao_05", "Quero um panorama da literatura sobre inversores.", True, ("inversor",)),
        ("revisao_06", "Quero a fundamentacao teorica da dissertacao.", True, ("FMEA",)),
        ("revisao_07", "Fale sobre FMEA.", False, ()),
        ("revisao_08", "Quais sao os resultados do AUC?", False, ()),
        # Extras com eh_revisao capturado pelo probe.
        ("revisao_09", "Quero a fundamentacao teorica completa.", True, ("FMEA",)),
        ("revisao_10", "Me da o referencial teorico da dissertacao.", True, ("autoencoder",)),
        ("revisao_11", "Faca uma sintese da literatura sobre PV.", True, ("Weibull",)),
        ("revisao_12", "Cite as referencias todas que temos.", True, ("RCM",)),
        ("revisao_13", "Cite os autores principais do tema.", True, ("inversor",)),
        ("revisao_14", "Quero um panorama geral da area.", True, ("FMEA",)),
        ("revisao_15", "Explique FMEA rapidamente.", False, ()),
        ("revisao_16", "Como interpreto o AUC?", False, ()),
        ("revisao_17", "Resuma o pipeline para mim.", False, ()),
        ("revisao_18", "Qual o proximo passo do projeto?", False, ()),
    ]
    for nome, pergunta, eh_rev, topicos in casos_revisao:
        casos.append(caso_expansao_revisao(nome, pergunta, eh_rev, topicos))

    # ── NOVO: leitura de anexos (unit) + prompt com bloco de anexo ──
    casos.extend(_casos_leitura_anexos())
    casos.extend(_casos_prompt_anexo())

    # ── NOVO: catalogo da literatura (inventario completo, anti-alucinacao) ──
    casos.extend(_casos_catalogo())

    # ── NOVO: experimentos de ML por artigo (roteamento + registry) ──
    casos.extend(_casos_experimentos())

    return casos


def gravar_memoria(resultados: list[dict], timestamp: str) -> int:
    # Reutiliza o modelo ja carregado pelo _rag_cache. Instanciar um SEGUNDO
    # SentenceTransformer aqui (com o primeiro ainda vivo no cache) provoca
    # access violation / segfault no Windows ao encodar — dois modelos torch
    # concorrendo pela mesma memoria nativa. Um unico modelo resolve.
    modelo, _ = _rag_cache()
    client = chromadb.PersistentClient(path=str(PASTA_CHROMADB))
    # Memória de AVALIAÇÃO, separada da memória de produção (sessoes_pv).
    colecao = client.get_or_create_collection(name=NOME_COLECAO_AVALIACOES)

    ids = []
    documentos = []
    metadados = []
    for item in resultados:
        status = "PASSOU" if item["ok"] else "FALHOU"
        doc = (
            f"# Avaliacao do agente - teste {item['indice']:03d}\n\n"
            f"- Data: {timestamp}\n"
            f"- Categoria: {item['categoria']}\n"
            f"- Caso: {item['nome']}\n"
            f"- Pergunta: {item['pergunta']}\n"
            f"- Resultado: {status}\n"
            f"- Aprendizado: {item['detalhe']}\n\n"
            "Memoria operacional: este teste calibra a politica de consulta a "
            "literatura, o roteamento de ferramentas, o tom conversacional, a "
            "formatacao de referencias e o prompt interno do Al IAdo PV."
        )
        ids.append(f"avaliacao_agente_100_{timestamp}_{item['indice']:03d}")
        documentos.append(doc)
        metadados.append({
            "tipo": "avaliacao_agente_100",
            "data": timestamp,
            "indice": str(item["indice"]),
            "categoria": item["categoria"],
            "ok": str(item["ok"]),
            "origem": "scripts/avaliar_agente_100.py",
        })

    embeddings = modelo.encode(documentos, show_progress_bar=True).tolist()
    upsert_em_lotes(colecao, ids, embeddings, documentos, metadados, tamanho_lote=100)
    verificados = colecao.get(ids=ids)
    return len(verificados.get("ids", []))


def gravar_relatorio(resultados: list[dict], memorias_gravadas: int, timestamp: str) -> Path:
    pasta = PASTA_NOTAS / "sessoes"
    pasta.mkdir(parents=True, exist_ok=True)
    caminho = pasta / f"{timestamp}_avaliacao_agente_100.md"

    total = len(resultados)
    falhas = [r for r in resultados if not r["ok"]]
    por_categoria: dict[str, list[dict]] = {}
    for item in resultados:
        por_categoria.setdefault(item["categoria"], []).append(item)

    linhas = [
        f"# Avaliacao tecnica do agente - {TOTAL_TESTES_ESPERADO} testes",
        "",
        f"- Data: {timestamp}",
        f"- Total de testes: {total}",
        f"- Passaram: {total - len(falhas)}",
        f"- Falharam: {len(falhas)}",
        f"- Memorias gravadas no ChromaDB: {memorias_gravadas}",
        "",
        "## Resultado por categoria",
        "",
    ]

    for categoria, itens in sorted(por_categoria.items()):
        ok = sum(1 for item in itens if item["ok"])
        linhas.append(f"- {categoria}: {ok}/{len(itens)}")

    linhas.extend(["", "## Casos", ""])
    linhas.append("| # | Categoria | Caso | Status | Detalhe |")
    linhas.append("|---:|---|---|---|---|")
    for item in resultados:
        status = "PASS" if item["ok"] else "FAIL"
        detalhe = str(item["detalhe"]).replace("|", "\\|").replace("\n", " ")
        linhas.append(
            f"| {item['indice']} | {item['categoria']} | {item['nome']} | {status} | {detalhe} |"
        )

    if falhas:
        linhas.extend(["", "## Falhas a corrigir", ""])
        for item in falhas:
            linhas.append(f"- {item['indice']:03d} {item['nome']}: {item['detalhe']}")
    else:
        linhas.extend(["", "## Parecer", ""])
        if memorias_gravadas:
            linhas.append(
                "A bateria passou integralmente. A colecao de avaliacao recebeu "
                "um documento por teste, registrando o comportamento esperado e "
                "o aprendizado operacional do agente."
            )
        else:
            linhas.append(
                "A bateria passou integralmente. Nenhuma memoria foi gravada; "
                "use --com-memoria apenas quando quiser registrar a avaliacao "
                "na colecao separada de testes."
            )

    caminho.write_text("\n".join(linhas) + "\n", encoding="utf-8")
    return caminho


TOTAL_TESTES_ESPERADO = 559


def main(gravar_memorias: bool = False) -> int:
    casos = montar_casos()
    if len(casos) != TOTAL_TESTES_ESPERADO:
        print(
            f"ERRO: a bateria precisa ter {TOTAL_TESTES_ESPERADO} testes, "
            f"mas tem {len(casos)}."
        )
        return 2

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    resultados = []

    for indice, caso in enumerate(casos, 1):
        try:
            ok, detalhe = caso.executar()
        except Exception as exc:
            ok = False
            detalhe = f"excecao={type(exc).__name__}: {exc}"

        resultados.append({
            "indice": indice,
            "nome": caso.nome,
            "categoria": caso.categoria,
            "pergunta": caso.pergunta,
            "ok": bool(ok),
            "detalhe": detalhe,
        })

    falhas = [r for r in resultados if not r["ok"]]

    if gravar_memorias:
        memorias_gravadas = gravar_memoria(resultados, timestamp)
        memorias_ok = memorias_gravadas == len(resultados)
        nota_memoria = str(memorias_gravadas)
    else:
        # Pula o encode pesado do torch (SentenceTransformer) que pode causar
        # access violation / segfault no Windows quando há outro modelo torch
        # vivo no mesmo processo. A verificação dos testes não depende disto.
        memorias_gravadas = 0
        memorias_ok = True
        nota_memoria = "(puladas: --sem-memoria)"

    relatorio = gravar_relatorio(resultados, memorias_gravadas, timestamp)

    print(f"Testes executados: {len(resultados)}")
    print(f"Passaram: {len(resultados) - len(falhas)}")
    print(f"Falharam: {len(falhas)}")
    print(f"Memorias gravadas: {nota_memoria}")
    print(f"Relatorio: {relatorio}")

    if falhas or not memorias_ok:
        return 1
    return 0


if __name__ == "__main__":
    import argparse

    _parser = argparse.ArgumentParser(
        description="Bateria determinística do Al IAdo PV."
    )
    # Por padrão NÃO grava memórias: avaliação não deve contaminar a memória de
    # produção (item 8.3) — e evita o segfault do torch no Windows. Use
    # --com-memoria para registrar explicitamente as memórias de teste.
    _parser.add_argument(
        "--com-memoria",
        action="store_true",
        help="Grava as memórias de teste no ChromaDB (NÃO é o padrão).",
    )
    _parser.add_argument(
        "--sem-memoria",
        action="store_true",
        help="(compat.) não grava memórias — já é o comportamento padrão.",
    )
    _args = _parser.parse_args()
    raise SystemExit(main(gravar_memorias=_args.com_memoria))
