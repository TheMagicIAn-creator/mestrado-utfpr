# -*- coding: utf-8 -*-
"""
============================================================================
 RENDERIZADOR DOCX  —  consome os blocos de `conteudo.py`
============================================================================
Layout corporativo enxuto, com logo no cabeçalho de todas as páginas,
marca d'água de fundo (imagem flutuante atrás do texto) e rodapé com
paginação. Toda a parte de OOXML (marca d'água, bordas, campos) está aqui.
============================================================================
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

import dados_projeto as D

COR = RGBColor.from_string(D.EMPRESA["cor_destaque"])
BRANCO = RGBColor.from_string("FFFFFF")
CINZA = RGBColor.from_string("595959")
ACENTO_HEX = D.EMPRESA["cor_destaque"]
ZEBRA_HEX = "F2F4F8"
CINZA_CLARO = "D9D9D9"


# ── helpers OOXML ──────────────────────────────────────────────────────────
def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_borders(table, cor="BFBFBF", sz=4):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), cor)
        borders.append(el)
    table._tbl.tblPr.append(borders)


def fixar_layout(table, larguras_cm):
    table.autofit = False
    table.allow_autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(larguras_cm[i])


def repetir_cabecalho(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true")
    trPr.append(th)


def celula(cell, texto, *, bold=False, cor=None, size=8.5, align="left", bg=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg:
        shade(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(str(texto))
    run.font.size = Pt(size); run.font.bold = bold
    if cor is not None:
        run.font.color.rgb = cor


def add_campo(paragraph, instr, default="1"):
    run = paragraph.add_run()
    for tipo, txt in (("begin", None), ("instr", instr),
                      ("separate", None), ("text", default), ("end", None)):
        if tipo == "instr":
            el = OxmlElement("w:instrText"); el.set(qn("xml:space"), "preserve")
            el.text = txt
        elif tipo == "text":
            el = OxmlElement("w:t"); el.text = txt
        else:
            el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), tipo)
        run._r.append(el)
    run.font.size = Pt(7.5); run.font.color.rgb = CINZA


# ── marca d'água (imagem flutuante atrás do texto) ─────────────────────────
def _anchor_xml(rId, cx, cy, nome="marca_dagua"):
    return (
        f'<w:drawing {nsdecls("w", "wp", "a", "pic", "r")}>'
        f'<wp:anchor behindDoc="1" distT="0" distB="0" distL="0" distR="0" '
        f'simplePos="0" locked="0" layoutInCell="1" allowOverlap="1" '
        f'relativeHeight="251654144">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:align>center</wp:align></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:align>center</wp:align></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
        f'<wp:docPr id="901" name="{nome}"/>'
        f'<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        f'<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic><pic:nvPicPr><pic:cNvPr id="901" name="{nome}"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic>'
        f'</a:graphicData></a:graphic></wp:anchor></w:drawing>'
    )


def add_marca_dagua(section, imagem, largura_cm=16.0):
    header = section.header
    header.is_linked_to_previous = False
    p = header.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(imagem, width=Cm(largura_cm))
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))
    extent = inline.find(qn("wp:extent"))
    cx, cy = extent.get("cx"), extent.get("cy")
    rId = inline.find(".//" + qn("a:blip")).get(qn("r:embed"))
    run._r.remove(drawing)
    run._r.append(parse_xml(_anchor_xml(rId, cx, cy)))


# ── cabeçalho / rodapé ─────────────────────────────────────────────────────
def _borda(cell, lado, sz="8", cor=ACENTO_HEX):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    b = OxmlElement(f"w:{lado}")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "0"); b.set(qn("w:color"), cor)
    borders.append(b); tcPr.append(borders)


def montar_cabecalho(section, logo, empresa, projeto):
    header = section.header
    header.is_linked_to_previous = False
    tbl = header.add_table(rows=1, cols=2, width=Cm(17))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    fixar_layout(tbl, [5.5, 11.5])
    c0 = tbl.rows[0].cells[0]
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if os.path.exists(logo):
        c0.paragraphs[0].add_run().add_picture(logo, width=Cm(4.6))
    c1 = tbl.rows[0].cells[1]
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for txt, size, bold, cor in [
        (empresa["nome"], 11, True, COR),
        (empresa["slogan"], 8, False, CINZA),
        (f'{projeto["codigo_doc"]}  ·  Rev. {projeto["revisao"]}', 8, False, CINZA),
    ]:
        p = c1.paragraphs[0] if txt == empresa["nome"] else c1.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt); r.font.size = Pt(size); r.font.bold = bold
        r.font.color.rgb = cor
    for cell in tbl.rows[0].cells:
        _borda(cell, "bottom")


def montar_rodape(section, empresa):
    footer = section.footer
    footer.is_linked_to_previous = False
    tbl = footer.add_table(rows=1, cols=3, width=Cm(17))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    fixar_layout(tbl, [6.0, 6.0, 5.0])
    for cell in tbl.rows[0].cells:
        _borda(cell, "top", "6", "BFBFBF")
    cL = tbl.rows[0].cells[0].paragraphs[0]
    cL.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cL.add_run(f'{empresa["cidade_uf"]}  ·  {empresa["cnpj"]}')
    r.font.size = Pt(7.5); r.font.color.rgb = CINZA
    cC = tbl.rows[0].cells[1].paragraphs[0]
    cC.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cC.add_run(empresa["contato"]); r.font.size = Pt(7.5); r.font.color.rgb = CINZA
    cR = tbl.rows[0].cells[2].paragraphs[0]
    cR.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = cR.add_run("Página "); r.font.size = Pt(7.5); r.font.color.rgb = CINZA
    add_campo(cR, " PAGE ", "1")
    r = cR.add_run(" de "); r.font.size = Pt(7.5); r.font.color.rgb = CINZA
    add_campo(cR, " NUMPAGES ", "1")


# ── blocos de conteúdo ─────────────────────────────────────────────────────
def _heading(doc, numero, titulo):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{numero}   {titulo}" if numero else titulo)
    r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = COR
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), ACENTO_HEX)
    pbdr.append(bottom); pPr.append(pbdr)


def _paragrafo(doc, texto, size=10, justify=True, space_after=6,
               bold=False, cor=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texto); r.font.size = Pt(size); r.font.bold = bold
    if cor is not None:
        r.font.color.rgb = cor


def _bullets(doc, itens, size=9.5):
    for it in itens:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it); r.font.size = Pt(size)


def _capa_titulo(doc, texto):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    fixar_layout(tbl, [17.0])
    cell = tbl.rows[0].cells[0]
    shade(cell, ACENTO_HEX)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(texto); r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = BRANCO
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _tabela_kv(doc, linhas, larg=(5.4, 11.6), size=9.5):
    t = doc.add_table(rows=len(linhas), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t, "D9D9D9", 4)
    fixar_layout(t, list(larg))
    for i, (k, v) in enumerate(linhas):
        celula(t.rows[i].cells[0], k, bold=True, size=size, bg=ZEBRA_HEX, cor=COR)
        celula(t.rows[i].cells[1], v, size=size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _tabela_previsao(doc, linhas, tot_ilum, tot_tug, tot_tomadas):
    cols = ["Ambiente", "Área (m²)", "Perím. (m)", "Iluminação (VA)",
            "Tomadas (un.)", "TUG (VA)"]
    larg = [4.6, 2.2, 2.4, 3.2, 2.4, 2.2]
    t = doc.add_table(rows=1, cols=len(cols))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t); fixar_layout(t, larg); repetir_cabecalho(t.rows[0])
    for i, c in enumerate(cols):
        celula(t.rows[0].cells[i], c, bold=True, cor=BRANCO, size=8.5,
               align="center", bg=ACENTO_HEX)
    for n, l in enumerate(linhas):
        row = t.add_row(); bg = ZEBRA_HEX if n % 2 else None
        celula(row.cells[0], l["ambiente"], size=8.5, bg=bg)
        celula(row.cells[1], f'{l["area"]:.1f}', size=8.5, align="center", bg=bg)
        celula(row.cells[2], f'{l["perimetro"]:.1f}', size=8.5, align="center", bg=bg)
        celula(row.cells[3], f'{l["ilum_va"]}', size=8.5, align="center", bg=bg)
        celula(row.cells[4], f'{l["n_tug"]}', size=8.5, align="center", bg=bg)
        celula(row.cells[5], f'{l["tug_va"]}', size=8.5, align="center", bg=bg)
    row = t.add_row()
    celula(row.cells[0], "TOTAL", bold=True, size=8.5, bg=CINZA_CLARO)
    for j in (1, 2):
        celula(row.cells[j], "", bg=CINZA_CLARO)
    celula(row.cells[3], f"{tot_ilum}", bold=True, size=8.5, align="center", bg=CINZA_CLARO)
    celula(row.cells[4], f"{tot_tomadas}", bold=True, size=8.5, align="center", bg=CINZA_CLARO)
    celula(row.cells[5], f"{tot_tug}", bold=True, size=8.5, align="center", bg=CINZA_CLARO)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _tabela_circuitos(doc, circuitos):
    cols = ["Circ.", "Descrição", "Sist./V", "Pot. (VA)", "IB (A)", "Fase (mm²)",
            "PE (mm²)", "Eletrod. (mm)", "Disjuntor", "DR (mA)", "ΔV (%)"]
    larg = [1.0, 4.6, 1.3, 1.4, 1.1, 1.2, 1.1, 1.4, 1.6, 1.2, 1.1]
    t = doc.add_table(rows=1, cols=len(cols))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t); fixar_layout(t, larg); repetir_cabecalho(t.rows[0])
    for i, c in enumerate(cols):
        celula(t.rows[0].cells[i], c, bold=True, cor=BRANCO, size=7.5,
               align="center", bg=ACENTO_HEX)
    for n, cir in enumerate(circuitos):
        row = t.add_row(); bg = ZEBRA_HEX if n % 2 else None
        vals = [
            (cir["id"], "center", True), (cir["descricao"], "left", False),
            (f'{cir["sistema"]} / {cir["tensao"]:.0f}', "center", False),
            (f'{cir["potencia_va"]:.0f}', "center", False),
            (f'{cir["ib"]:.1f}', "center", False),
            (f'{cir["secao"]:.1f}', "center", False),
            (f'{cir["pe"]:.1f}', "center", False),
            (f'{cir["eletroduto"]}', "center", False),
            (f'{cir["disjuntor"]} (C)', "center", False),
            ("30" if cir["dr"] else "–", "center", False),
            (f'{cir["dv"]:.2f}', "center", False),
        ]
        for j, (v, al, bd) in enumerate(vals):
            celula(row.cells[j], v, size=7.5, align=al, bg=bg, bold=bd)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _assinatura(doc, linhas):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_" * 48)
    for texto, bold in linhas:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_after = Pt(0)
        r = pp.add_run(texto); r.font.size = Pt(10); r.font.bold = bold


# ── orquestração ───────────────────────────────────────────────────────────
def gerar_docx(ctx, blocos, caminho):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string("262626")

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(0.8)

    montar_cabecalho(sec, ctx["logo"], D.EMPRESA, D.PROJETO)
    montar_rodape(sec, D.EMPRESA)
    add_marca_dagua(sec, ctx["marca"], largura_cm=16.0)

    for b in blocos:
        tp = b["tipo"]
        if tp == "capa_titulo":
            _capa_titulo(doc, b["texto"])
        elif tp == "capa_ident":
            _tabela_kv(doc, b["linhas"], larg=(4.8, 12.2))
        elif tp == "secao":
            _heading(doc, b["num"], b["titulo"])
        elif tp == "subsecao":
            _paragrafo(doc, b["titulo"], bold=True, justify=False, space_after=2)
        elif tp == "paragrafo":
            _paragrafo(doc, b["texto"], size=b.get("size", 10))
        elif tp == "nota":
            _paragrafo(doc, b["texto"], size=8, space_after=6, cor=CINZA)
        elif tp == "legenda":
            _paragrafo(doc, b["texto"], size=8, space_after=6, cor=CINZA)
        elif tp == "bullets":
            _bullets(doc, b["itens"], size=b.get("size", 9.5))
        elif tp == "tabela_kv":
            _tabela_kv(doc, b["linhas"])
        elif tp == "tabela_previsao":
            _tabela_previsao(doc, b["linhas"], b["tot_ilum"], b["tot_tug"], b["tot_tomadas"])
        elif tp == "tabela_circuitos":
            _tabela_circuitos(doc, b["circuitos"])
        elif tp == "espaco":
            doc.add_paragraph().paragraph_format.space_after = Pt(b.get("pt", 12))
        elif tp == "assinatura":
            _assinatura(doc, b["linhas"])

    cp = doc.core_properties
    cp.title = D.PROJETO["titulo"]
    cp.author = D.PROJETO["responsavel"]
    cp.subject = "Memorial Descritivo de Instalação Elétrica – ABNT NBR 5410/5419"
    cp.category = "Projeto Elétrico"

    # faz o Word recalcular os campos (Página X de Y) ao abrir o arquivo
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        uf = OxmlElement("w:updateFields")
        uf.set(qn("w:val"), "true")
        settings.append(uf)

    doc.save(caminho)
    return caminho
